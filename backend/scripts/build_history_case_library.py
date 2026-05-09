#!/usr/bin/env python3
"""Build a searchable local library from historical bid documents.

The script scans historical folders named like ``2024年招投标-技术1`` and imports
all project folders under them. It stores structured metadata, full-text search
content, and PageIndex-compatible tree JSON in ``backend/data``.
"""
from __future__ import annotations

import argparse
import hashlib
import json
import re
import sqlite3
import subprocess
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterable

import PyPDF2
import docx


REPO_ROOT = Path(__file__).resolve().parents[2]
DEFAULT_HISTORY_ROOT = REPO_ROOT
DEFAULT_DB_PATH = REPO_ROOT / "backend" / "data" / "history_cases.sqlite3"
DEFAULT_ARTIFACT_ROOT = REPO_ROOT / "backend" / "data" / "history_cases"
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc"}
RESULT_TAGS = ("中标", "未中", "流标", "废标", "放弃", "陪标", "待定", "延期")
DOMAIN_RULES = [
    ("石油", "石油", ("石油", "中石油", "中国石油", "中石化", "中国石化", "石化", "成品油", "油品", "销售公司", "航油")),
    ("燃气", "天然气/LNG", ("燃气", "天然气", "LNG", "液化天然气", "加气", "气化", "气电", "CNG")),
    ("化工", "化工/煤化工", ("化工", "煤化工", "炼化", "纯苯", "合成氨", "甲基环己烷", "焦油", "苯", "DMC", "硝酸", "化肥", "水溶肥")),
    ("电力新能源", "光伏/充电/电力", ("光伏", "新能源", "充电", "电力", "国网", "66kV", "风电", "储能", "光火储", "电气")),
    ("氢能", "制氢/加氢", ("氢", "制氢", "加氢", "输氢", "储氢")),
    ("油库加油站", "油库/加油站", ("油库", "加油站", "加能站", "综合能源站", "油罐", "储罐", "罐区", "换底", "油料")),
    ("管道", "管道/输气输油", ("管道", "输气", "输油", "西气东输", "场站", "增压", "管廊", "管线", "防爆区域")),
    ("军队油料", "部队/某部油料", ("某部", "部队", "油料", "单位油库", "军", "营区")),
    ("工程设计服务", "设计/咨询/框架", ("设计", "勘察", "可研", "初设", "施工图", "咨询", "框架", "技术服务", "工程服务", "建模")),
    ("基础设施", "市政/交通/园区", ("服务区", "高速", "有轨电车", "园区", "城市燃气", "供水", "污水", "供暖", "建筑", "市政")),
    ("信息化数字化", "建模/数字化/平台", ("三维", "建模", "数字化", "平台", "物联网", "远程诊断", "监视")),
]
DOMAIN_PRIORITY = {
    "石油": 100,
    "燃气": 98,
    "化工": 96,
    "电力新能源": 94,
    "氢能": 92,
    "油库加油站": 90,
    "管道": 88,
    "军队油料": 86,
    "信息化数字化": 72,
    "基础设施": 65,
    "工程设计服务": 30,
    "其他": 0,
}


@dataclass(frozen=True)
class HistoricalProject:
    project_id: str
    year: str
    batch: str
    sequence: str
    result: str
    subject: str
    title: str
    folder_name: str
    source_path: Path


@dataclass(frozen=True)
class DomainMatch:
    domain: str
    subdomain: str
    confidence: float
    keywords: tuple[str, ...]


def now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def stable_id(prefix: str, value: str) -> str:
    digest = hashlib.sha1(value.encode("utf-8")).hexdigest()[:16]
    return f"{prefix}-{digest}"


def classify_domains(*parts: str) -> list[DomainMatch]:
    """Classify project/domain with transparent keyword rules."""
    text = "\n".join(part for part in parts if part)
    if not text.strip():
        return [DomainMatch("其他", "未分类", 0.2, tuple())]

    matches: list[DomainMatch] = []
    lower_text = text.lower()
    for domain, subdomain, keywords in DOMAIN_RULES:
        hit_keywords = []
        score = 0
        for keyword in keywords:
            keyword_lower = keyword.lower()
            count = lower_text.count(keyword_lower)
            if count:
                hit_keywords.append(keyword)
                score += count
        if hit_keywords:
            confidence = min(0.99, 0.45 + 0.08 * len(hit_keywords) + 0.03 * min(score, 8))
            matches.append(
                DomainMatch(
                    domain=domain,
                    subdomain=subdomain,
                    confidence=round(confidence, 2),
                    keywords=tuple(hit_keywords),
                )
            )

    if not matches:
        return [DomainMatch("其他", "未分类", 0.2, tuple())]
    matches.sort(key=lambda item: (-(DOMAIN_PRIORITY.get(item.domain, 0)), -item.confidence, item.domain))
    return matches


def parse_project_folder(parent_name: str, project_path: Path) -> HistoricalProject:
    folder_match = re.match(r"(20\d{2})年招投标-技术(\d+)$", parent_name)
    year = folder_match.group(1) if folder_match else ""
    batch = f"技术{folder_match.group(2)}" if folder_match else ""

    raw_name = project_path.name.strip()
    seq_match = re.match(r"^(\d+(?:\.\d+)?)\s*(.*)$", raw_name)
    sequence = seq_match.group(1) if seq_match else ""
    rest = seq_match.group(2).strip() if seq_match else raw_name

    tag = ""
    title = rest
    tag_match = re.match(r"^【([^】]+)】\s*(.*)$", rest)
    if tag_match:
        tag = tag_match.group(1).strip()
        title = tag_match.group(2).strip()

    result = "、".join([item for item in RESULT_TAGS if item in raw_name]) or "未标明"
    subject = tag
    for item in RESULT_TAGS:
        subject = subject.replace(item, "")
    subject = re.sub(r"[-_—－+、，,\.\s]+$", "", subject)
    subject = re.sub(r"^[-_—－+、，,\.\s]+", "", subject)
    subject = re.sub(r"[-_—－]+", "-", subject).strip() or tag or "未标明"

    return HistoricalProject(
        project_id=stable_id("case", str(project_path.resolve())),
        year=year,
        batch=batch,
        sequence=sequence,
        result=result,
        subject=subject,
        title=title,
        folder_name=raw_name,
        source_path=project_path.resolve(),
    )


def iter_historical_projects(history_root: Path, winning_only: bool = False) -> Iterable[HistoricalProject]:
    for outer in sorted(history_root.iterdir(), key=lambda item: item.name):
        if not outer.is_dir() or not re.match(r"20\d{2}年招投标-技术\d+$", outer.name):
            continue
        inner = outer / outer.name
        if not inner.is_dir():
            continue
        for project_path in sorted(inner.iterdir(), key=lambda item: item.name):
            if project_path.is_dir() and (not winning_only or "中标" in project_path.name):
                yield parse_project_folder(outer.name, project_path)


def iter_documents(project: HistoricalProject) -> Iterable[Path]:
    for path in sorted(project.source_path.rglob("*"), key=lambda item: str(item)):
        if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS:
            yield path.resolve()


def heading_level(paragraph: docx.text.paragraph.Paragraph) -> int:
    style_name = (paragraph.style.name if paragraph.style else "") or ""
    match = re.search(r"heading\s*(\d+)|标题\s*(\d+)", style_name, re.IGNORECASE)
    if match:
        return max(1, min(6, int(match.group(1) or match.group(2))))

    text = paragraph.text.strip()
    if not text or len(text) > 80:
        return 0
    if re.match(r"^第[一二三四五六七八九十百\d]+[章节篇部分]", text):
        return 1
    if re.match(r"^[一二三四五六七八九十]+[、.．]", text):
        return 2
    if re.match(r"^（[一二三四五六七八九十]+）", text):
        return 3
    if re.match(r"^\d+(?:\.\d+)*[、.．]\s*", text):
        return min(4, text.count(".") + text.count("．") + 2)
    return 0


def extract_docx_markdown(path: Path) -> str:
    document = docx.Document(str(path))
    lines: list[str] = [f"# {path.stem}", ""]

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        level = heading_level(paragraph)
        if level:
            lines.append(f"{'#' * level} {text}")
        else:
            lines.append(text)
        lines.append("")

    for index, table in enumerate(document.tables, 1):
        lines.extend([f"## 表格 {index}", ""])
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                lines.append("| " + " | ".join(cells) + " |")
        lines.append("")

    return "\n".join(lines).strip() + "\n"


def extract_pdf_markdown(path: Path) -> str:
    lines: list[str] = [f"# {path.stem}", ""]
    with path.open("rb") as file:
        reader = PyPDF2.PdfReader(file)
        for index, page in enumerate(reader.pages, 1):
            text = page.extract_text() or ""
            text = re.sub(r"\n{3,}", "\n\n", text).strip()
            lines.extend([f"## 第 {index} 页", "", text or "（本页未提取到文本）", ""])
    return "\n".join(lines).strip() + "\n"


def extract_doc_markdown(path: Path) -> str:
    """Extract legacy .doc files through macOS textutil."""
    proc = subprocess.run(
        ["textutil", "-convert", "txt", "-stdout", str(path)],
        capture_output=True,
        text=True,
        timeout=120,
        check=False,
    )
    if proc.returncode != 0:
        error = (proc.stderr or proc.stdout or "textutil conversion failed").strip()
        sibling_pdf = path.with_suffix(".pdf")
        if sibling_pdf.exists():
            pdf_markdown = extract_pdf_markdown(sibling_pdf)
            return (
                f"# {path.stem}\n\n"
                f"（旧版 DOC 通过 textutil 解析失败，已使用同名 PDF 作为检索文本来源：{sibling_pdf.name}）\n\n"
                f"{pdf_markdown}\n"
            )
        raise RuntimeError(error)

    body = proc.stdout.strip()
    if not body:
        sibling_pdf = path.with_suffix(".pdf")
        if sibling_pdf.exists():
            pdf_markdown = extract_pdf_markdown(sibling_pdf)
            return (
                f"# {path.stem}\n\n"
                f"（旧版 DOC 未提取到文本，已使用同名 PDF 作为检索文本来源：{sibling_pdf.name}）\n\n"
                f"{pdf_markdown}\n"
            )
        raise RuntimeError("textutil 未提取到文本")
    body = re.sub(r"\n{3,}", "\n\n", body)
    return f"# {path.stem}\n\n{body}\n"


def markdown_to_pageindex_tree(markdown: str, doc_name: str) -> dict:
    """Create a PageIndex Markdown-mode compatible tree without LLM summaries."""
    lines = markdown.splitlines()
    headers: list[dict] = []
    for line_number, line in enumerate(lines, 1):
        match = re.match(r"^(#{1,6})\s+(.+?)\s*$", line)
        if match:
            headers.append(
                {
                    "level": len(match.group(1)),
                    "title": match.group(2).strip(),
                    "line_num": line_number,
                }
            )

    if not headers:
        headers.append({"level": 1, "title": doc_name, "line_num": 1})

    for index, header in enumerate(headers):
        start = header["line_num"] - 1
        end = headers[index + 1]["line_num"] - 1 if index + 1 < len(headers) else len(lines)
        header["text"] = "\n".join(lines[start:end]).strip()

    root_nodes: list[dict] = []
    stack: list[tuple[int, dict]] = []
    for index, header in enumerate(headers, 1):
        node = {
            "title": header["title"],
            "node_id": str(index).zfill(4),
            "line_num": header["line_num"],
            "text": header["text"],
            "nodes": [],
        }
        while stack and stack[-1][0] >= header["level"]:
            stack.pop()
        if stack:
            stack[-1][1]["nodes"].append(node)
        else:
            root_nodes.append(node)
        stack.append((header["level"], node))

    return {
        "doc_name": doc_name,
        "line_count": len(lines),
        "structure": root_nodes,
    }


def walk_nodes(nodes: list[dict], parent_id: str | None = None, level: int = 1) -> Iterable[dict]:
    for node in nodes:
        current = {
            "node_id": node.get("node_id", ""),
            "parent_node_id": parent_id,
            "title": node.get("title", ""),
            "line_num": int(node.get("line_num") or 0),
            "level": level,
            "text": node.get("text", ""),
        }
        yield current
        yield from walk_nodes(node.get("nodes") or [], current["node_id"], level + 1)


def connect(db_path: Path) -> sqlite3.Connection:
    db_path.parent.mkdir(parents=True, exist_ok=True)
    conn = sqlite3.connect(str(db_path))
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA journal_mode=WAL")
    conn.execute("PRAGMA foreign_keys=ON")
    ensure_schema(conn)
    return conn


def ensure_schema(conn: sqlite3.Connection) -> None:
    conn.executescript(
        """
        CREATE TABLE IF NOT EXISTS case_projects (
            id TEXT PRIMARY KEY,
            year TEXT NOT NULL,
            batch TEXT NOT NULL,
            sequence TEXT NOT NULL,
            result TEXT NOT NULL,
            subject TEXT NOT NULL,
            primary_domain TEXT NOT NULL DEFAULT '其他',
            primary_subdomain TEXT NOT NULL DEFAULT '未分类',
            domain_confidence REAL NOT NULL DEFAULT 0,
            domain_keywords TEXT NOT NULL DEFAULT '[]',
            title TEXT NOT NULL,
            folder_name TEXT NOT NULL,
            source_path TEXT NOT NULL UNIQUE,
            imported_at TEXT NOT NULL
        );

        CREATE TABLE IF NOT EXISTS case_project_domains (
            project_id TEXT NOT NULL REFERENCES case_projects(id) ON DELETE CASCADE,
            domain TEXT NOT NULL,
            subdomain TEXT NOT NULL,
            confidence REAL NOT NULL,
            keywords TEXT NOT NULL DEFAULT '[]',
            PRIMARY KEY(project_id, domain, subdomain)
        );

        CREATE TABLE IF NOT EXISTS case_documents (
            id TEXT PRIMARY KEY,
            project_id TEXT NOT NULL REFERENCES case_projects(id) ON DELETE CASCADE,
            file_name TEXT NOT NULL,
            extension TEXT NOT NULL,
            source_path TEXT NOT NULL UNIQUE,
            markdown_path TEXT NOT NULL,
            pageindex_tree_path TEXT NOT NULL,
            text_chars INTEGER NOT NULL DEFAULT 0,
            node_count INTEGER NOT NULL DEFAULT 0,
            status TEXT NOT NULL,
            error TEXT NOT NULL DEFAULT '',
            indexed_at TEXT NOT NULL
        );

        CREATE TABLE IF NOT EXISTS pageindex_nodes (
            id TEXT PRIMARY KEY,
            document_id TEXT NOT NULL REFERENCES case_documents(id) ON DELETE CASCADE,
            node_id TEXT NOT NULL,
            parent_node_id TEXT,
            title TEXT NOT NULL,
            line_num INTEGER NOT NULL,
            level INTEGER NOT NULL,
            text TEXT NOT NULL
        );

        CREATE VIRTUAL TABLE IF NOT EXISTS history_case_fts USING fts5(
            project_id UNINDEXED,
            document_id UNINDEXED,
            project_title,
            file_name,
            body,
            tokenize='unicode61'
        );
        """
    )
    ensure_column(conn, "case_projects", "primary_domain", "TEXT NOT NULL DEFAULT '其他'")
    ensure_column(conn, "case_projects", "primary_subdomain", "TEXT NOT NULL DEFAULT '未分类'")
    ensure_column(conn, "case_projects", "domain_confidence", "REAL NOT NULL DEFAULT 0")
    ensure_column(conn, "case_projects", "domain_keywords", "TEXT NOT NULL DEFAULT '[]'")
    conn.commit()


def ensure_column(conn: sqlite3.Connection, table: str, column: str, definition: str) -> None:
    columns = {row["name"] for row in conn.execute(f"PRAGMA table_info({table})").fetchall()}
    if column not in columns:
        conn.execute(f"ALTER TABLE {table} ADD COLUMN {column} {definition}")


def reset_library(conn: sqlite3.Connection) -> None:
    conn.executescript(
        """
        DELETE FROM history_case_fts;
        DELETE FROM pageindex_nodes;
        DELETE FROM case_documents;
        DELETE FROM case_project_domains;
        DELETE FROM case_projects;
        """
    )
    conn.commit()


def write_text(path: Path, content: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding="utf-8")


def write_json(path: Path, payload: dict) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def import_project(conn: sqlite3.Connection, project: HistoricalProject) -> None:
    domain_matches = classify_domains(project.title, project.folder_name, project.subject)
    primary = domain_matches[0]
    conn.execute(
        """
        INSERT INTO case_projects(
            id, year, batch, sequence, result, subject, primary_domain, primary_subdomain,
            domain_confidence, domain_keywords, title, folder_name, source_path, imported_at
        ) VALUES(?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ON CONFLICT(source_path) DO UPDATE SET
            year=excluded.year,
            batch=excluded.batch,
            sequence=excluded.sequence,
            result=excluded.result,
            subject=excluded.subject,
            primary_domain=excluded.primary_domain,
            primary_subdomain=excluded.primary_subdomain,
            domain_confidence=excluded.domain_confidence,
            domain_keywords=excluded.domain_keywords,
            title=excluded.title,
            folder_name=excluded.folder_name,
            imported_at=excluded.imported_at
        """,
        (
            project.project_id,
            project.year,
            project.batch,
            project.sequence,
            project.result,
            project.subject,
            primary.domain,
            primary.subdomain,
            primary.confidence,
            json.dumps(primary.keywords, ensure_ascii=False),
            project.title,
            project.folder_name,
            str(project.source_path),
            now_iso(),
        ),
    )
    conn.execute("DELETE FROM case_project_domains WHERE project_id=?", (project.project_id,))
    for match in domain_matches:
        conn.execute(
            """
            INSERT INTO case_project_domains(project_id, domain, subdomain, confidence, keywords)
            VALUES(?, ?, ?, ?, ?)
            """,
            (
                project.project_id,
                match.domain,
                match.subdomain,
                match.confidence,
                json.dumps(match.keywords, ensure_ascii=False),
            ),
        )


def import_document(conn: sqlite3.Connection, project: HistoricalProject, doc_path: Path, artifact_root: Path) -> tuple[str, str]:
    doc_id = stable_id("doc", str(doc_path.resolve()))
    relative_base = Path(project.year) / project.batch / f"{project.sequence}-{doc_id.removeprefix('doc-')}"
    markdown_path = artifact_root / "markdown" / relative_base / f"{doc_path.stem}.md"
    tree_path = artifact_root / "pageindex_trees" / relative_base / f"{doc_path.stem}.json"
    status = "indexed"
    error = ""

    try:
        if doc_path.suffix.lower() == ".docx":
            markdown = extract_docx_markdown(doc_path)
        elif doc_path.suffix.lower() == ".doc":
            markdown = extract_doc_markdown(doc_path)
        elif doc_path.suffix.lower() == ".pdf":
            markdown = extract_pdf_markdown(doc_path)
        else:
            raise ValueError(f"unsupported extension: {doc_path.suffix}")

        tree = markdown_to_pageindex_tree(markdown, doc_path.stem)
        nodes = list(walk_nodes(tree["structure"]))
        write_text(markdown_path, markdown)
        write_json(tree_path, tree)
    except Exception as exc:
        markdown = ""
        tree = {"doc_name": doc_path.stem, "line_count": 0, "structure": []}
        nodes = []
        status = "failed"
        error = str(exc)

    conn.execute("DELETE FROM pageindex_nodes WHERE document_id=?", (doc_id,))
    conn.execute("DELETE FROM history_case_fts WHERE document_id=?", (doc_id,))
    conn.execute(
        """
        INSERT INTO case_documents(
            id, project_id, file_name, extension, source_path, markdown_path, pageindex_tree_path,
            text_chars, node_count, status, error, indexed_at
        ) VALUES(?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ON CONFLICT(source_path) DO UPDATE SET
            project_id=excluded.project_id,
            file_name=excluded.file_name,
            extension=excluded.extension,
            markdown_path=excluded.markdown_path,
            pageindex_tree_path=excluded.pageindex_tree_path,
            text_chars=excluded.text_chars,
            node_count=excluded.node_count,
            status=excluded.status,
            error=excluded.error,
            indexed_at=excluded.indexed_at
        """,
        (
            doc_id,
            project.project_id,
            doc_path.name,
            doc_path.suffix.lower(),
            str(doc_path),
            str(markdown_path),
            str(tree_path),
            len(markdown),
            len(nodes),
            status,
            error,
            now_iso(),
        ),
    )

    for node in nodes:
        conn.execute(
            """
            INSERT INTO pageindex_nodes(
                id, document_id, node_id, parent_node_id, title, line_num, level, text
            ) VALUES(?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                f"{doc_id}:{node['node_id']}",
                doc_id,
                node["node_id"],
                node["parent_node_id"],
                node["title"],
                node["line_num"],
                node["level"],
                node["text"],
            ),
        )

    if markdown:
        conn.execute(
            """
            INSERT INTO history_case_fts(project_id, document_id, project_title, file_name, body)
            VALUES(?, ?, ?, ?, ?)
            """,
            (project.project_id, doc_id, project.title, doc_path.name, markdown),
        )

    return status, error


def update_project_domain_from_documents(conn: sqlite3.Connection, project: HistoricalProject) -> None:
    rows = conn.execute(
        """
        SELECT d.file_name, f.body
        FROM case_documents d
        LEFT JOIN history_case_fts f ON f.document_id = d.id
        WHERE d.project_id=?
        """,
        (project.project_id,),
    ).fetchall()
    file_names = "\n".join(str(row["file_name"] or "") for row in rows)
    body_preview = "\n".join(str(row["body"] or "")[:6000] for row in rows)
    primary_matches = classify_domains(project.title, project.folder_name, project.subject, file_names)
    domain_matches = classify_domains(project.title, project.folder_name, project.subject, file_names, body_preview)
    primary = primary_matches[0]
    conn.execute(
        """
        UPDATE case_projects
        SET primary_domain=?, primary_subdomain=?, domain_confidence=?, domain_keywords=?
        WHERE id=?
        """,
        (
            primary.domain,
            primary.subdomain,
            primary.confidence,
            json.dumps(primary.keywords, ensure_ascii=False),
            project.project_id,
        ),
    )
    conn.execute("DELETE FROM case_project_domains WHERE project_id=?", (project.project_id,))
    for match in domain_matches:
        conn.execute(
            """
            INSERT INTO case_project_domains(project_id, domain, subdomain, confidence, keywords)
            VALUES(?, ?, ?, ?, ?)
            """,
            (
                project.project_id,
                match.domain,
                match.subdomain,
                match.confidence,
                json.dumps(match.keywords, ensure_ascii=False),
            ),
        )


def build_library(history_root: Path, db_path: Path, artifact_root: Path, reset: bool, winning_only: bool = False) -> dict:
    projects = list(iter_historical_projects(history_root, winning_only=winning_only))
    with connect(db_path) as conn:
        if reset:
            reset_library(conn)

        total_docs = 0
        indexed_docs = 0
        failed_docs = 0
        for project in projects:
            import_project(conn, project)
            for doc_path in iter_documents(project):
                total_docs += 1
                status, _ = import_document(conn, project, doc_path, artifact_root)
                if status == "indexed":
                    indexed_docs += 1
                else:
                    failed_docs += 1
            update_project_domain_from_documents(conn, project)
            conn.commit()

    return {
        "scope": "winning_only" if winning_only else "all_history",
        "projects": len(projects),
        "documents": total_docs,
        "indexed_documents": indexed_docs,
        "failed_documents": failed_docs,
        "db_path": str(db_path),
        "artifact_root": str(artifact_root),
    }


def search_library(db_path: Path, query: str, limit: int) -> list[dict]:
    with connect(db_path) as conn:
        rows = conn.execute(
            """
            SELECT
                bm25(history_case_fts) AS rank,
                p.year,
                p.batch,
                p.sequence,
                p.subject,
                p.result,
                p.title AS project_title,
                d.file_name,
                d.source_path,
                snippet(history_case_fts, 4, '[', ']', '...', 18) AS snippet
            FROM history_case_fts
            JOIN case_projects p ON p.id = history_case_fts.project_id
            JOIN case_documents d ON d.id = history_case_fts.document_id
            WHERE history_case_fts MATCH ?
            ORDER BY rank
            LIMIT ?
            """,
            (query, limit),
        ).fetchall()
        results = [dict(row) for row in rows]
        seen = {row["source_path"] for row in results}
        if len(results) >= limit:
            return results

        like = f"%{query}%"
        fallback_rows = conn.execute(
            """
            SELECT
                0.0 AS rank,
                p.year,
                p.batch,
                p.sequence,
                p.subject,
                p.result,
                p.title AS project_title,
                d.file_name,
                d.source_path,
                history_case_fts.body AS body
            FROM history_case_fts
            JOIN case_projects p ON p.id = history_case_fts.project_id
            JOIN case_documents d ON d.id = history_case_fts.document_id
            WHERE
                p.title LIKE ?
                OR p.subject LIKE ?
                OR p.result LIKE ?
                OR p.folder_name LIKE ?
                OR d.file_name LIKE ?
                OR history_case_fts.body LIKE ?
            ORDER BY p.year, p.batch, CAST(p.sequence AS REAL), p.sequence, d.file_name
            LIMIT ?
            """,
            (like, like, like, like, like, like, limit * 3),
        ).fetchall()
        for row in fallback_rows:
            if row["source_path"] in seen:
                continue
            record = dict(row)
            body = str(record.pop("body") or "")
            index = body.find(query)
            record["snippet"] = body[max(0, index - 70): index + len(query) + 110] if index >= 0 else body[:180]
            results.append(record)
            seen.add(row["source_path"])
            if len(results) >= limit:
                break
        return results


def main() -> None:
    parser = argparse.ArgumentParser(description="Build/search historical bid case library.")
    parser.add_argument("--history-root", type=Path, default=DEFAULT_HISTORY_ROOT)
    parser.add_argument("--db-path", type=Path, default=DEFAULT_DB_PATH)
    parser.add_argument("--artifact-root", type=Path, default=DEFAULT_ARTIFACT_ROOT)
    parser.add_argument("--reset", action="store_true", help="Clear existing imported case data before building.")
    parser.add_argument("--winning-only", action="store_true", help="Import only projects whose folder name contains 中标.")
    parser.add_argument("--search", type=str, default="", help="Search the built FTS library instead of rebuilding.")
    parser.add_argument("--limit", type=int, default=10)
    args = parser.parse_args()

    if args.search:
        print(json.dumps(search_library(args.db_path, args.search, args.limit), ensure_ascii=False, indent=2))
        return

    summary = build_library(
        history_root=args.history_root.resolve(),
        db_path=args.db_path.resolve(),
        artifact_root=args.artifact_root.resolve(),
        reset=args.reset,
        winning_only=args.winning_only,
    )
    print(json.dumps(summary, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
