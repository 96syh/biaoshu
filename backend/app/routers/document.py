"""文档处理相关API路由"""
from fastapi import APIRouter, UploadFile, File, HTTPException
from fastapi.responses import JSONResponse, StreamingResponse
from ..models.schemas import (
    AnalysisReportRequest,
    AnalysisTaskControlRequest,
    AnalysisTaskControlResponse,
    AnalysisRequest,
    AnalysisType,
    ComplianceReviewRequest,
    ConsistencyRevisionRequest,
    DocumentBlocksPlanRequest,
    FileUploadResponse,
    WordExportRequest,
)
from ..services.file_service import FileService
from ..services.openai_service import OpenAIService
from ..utils.config_manager import config_manager
from ..utils.provider_registry import get_provider_auth_error
from ..utils.sse import sse_response
import json
import io
import re
import docx
import asyncio
import time
import uuid
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from urllib.parse import quote
from pathlib import Path

router = APIRouter(prefix="/api/document", tags=["文档处理"])


class AnalysisTaskState:
    """Keep server-side control state for one standard-analysis SSE task."""

    def __init__(self, task_id: str):
        self.task_id = task_id
        self.pause_event = asyncio.Event()
        self.pause_event.set()
        self.cancelled = False
        self.status = "running"
        self.current_step = 0


ANALYSIS_TASKS: dict[str, AnalysisTaskState] = {}


def _analysis_metric_count(report: dict, keys: tuple[str, ...]) -> int:
    return sum(len(report.get(key) or []) for key in keys)


def _analysis_stream_payload(
    task_state: AnalysisTaskState,
    step_index: int,
    detail: str,
    percent: int,
    status: str = "running",
    **extra,
) -> str:
    task_state.current_step = step_index
    task_state.status = status
    payload = {
        "chunk": "",
        "task_id": task_state.task_id,
        "step_index": step_index,
        "detail": detail,
        "percent": max(0, min(100, int(percent))),
        "status": status,
        **extra,
    }
    return f"data: {json.dumps(payload, ensure_ascii=False)}\n\n"


def sanitize_docx_filename(filename: str) -> str:
    """清理保存到本地文件系统的 docx 文件名。"""
    safe = re.sub(r'[\\/:*?"<>|\r\n]+', "_", filename or "标书文档.docx").strip()
    if not safe.lower().endswith(".docx"):
        safe = f"{safe}.docx"
    return safe or "标书文档.docx"


def set_run_font_simsun(run: docx.text.run.Run) -> None:
    """统一将 run 字体设置为宋体（包含 EastAsia 字体设置）"""
    run.font.name = "宋体"
    r = run._element.rPr
    if r is not None and r.rFonts is not None:
        r.rFonts.set(qn("w:eastAsia"), "宋体")


def set_paragraph_font_simsun(paragraph: docx.text.paragraph.Paragraph) -> None:
    """将段落内所有 runs 字体设置为宋体"""
    for run in paragraph.runs:
        set_run_font_simsun(run)


def add_field(paragraph: docx.text.paragraph.Paragraph, instruction: str) -> None:
    """Add a Word field that can be updated inside Word/WPS."""
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "请在 Word 中更新域"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    run = paragraph.add_run()
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(placeholder)
    run._r.append(end)
    set_run_font_simsun(run)


def enable_update_fields_on_open(doc: docx.Document) -> None:
    """Ask Word to update TOC/page-number fields when opening the document."""
    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
    settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def _style_value(style_profile: dict, key: str, default: str) -> str:
    value = style_profile.get(key)
    return str(value).strip() if value else default


def _cm_value(value: str, default: float) -> float:
    text = str(value or "").strip().lower().replace("厘米", "cm")
    try:
        if text.endswith("cm"):
            return float(text[:-2])
        if text.endswith("mm"):
            return float(text[:-2]) / 10
        return float(text)
    except Exception:
        return default


def _pt_value(value: str, default: float) -> float:
    text = str(value or "").strip().lower().replace("小四", "12pt").replace("五号", "10.5pt")
    try:
        if text.endswith("pt"):
            return float(text[:-2])
        return float(text)
    except Exception:
        return default


def configure_word_style(doc: docx.Document, reference_profile: dict | None = None) -> None:
    """Apply mature-sample-like Word page and font settings."""
    profile = reference_profile or {}
    style_profile = profile.get("word_style_profile") if isinstance(profile, dict) else {}
    if not isinstance(style_profile, dict):
        style_profile = {}

    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(_cm_value(_style_value(style_profile, "margin_top", "2.2cm"), 2.2))
        section.bottom_margin = Cm(_cm_value(_style_value(style_profile, "margin_bottom", "2.2cm"), 2.2))
        section.left_margin = Cm(_cm_value(_style_value(style_profile, "margin_left", "2.7cm"), 2.7))
        section.right_margin = Cm(_cm_value(_style_value(style_profile, "margin_right", "2.2cm"), 2.2))
        section.header_distance = Cm(1.5)
        section.footer_distance = Cm(1.75)

    body_font = _style_value(style_profile, "body_font_family", "宋体").split(",")[0].strip().strip('"')
    heading_font = _style_value(style_profile, "heading_font_family", "黑体").split(",")[0].strip().strip('"')
    body_size = _pt_value(_style_value(style_profile, "body_font_size", "10.5pt"), 10.5)

    def set_style_font(style_name: str, font_name: str, size_pt: float, bold: bool | None = None) -> None:
        style = doc.styles[style_name]
        style.font.name = font_name
        style.font.size = Pt(size_pt)
        if bold is not None:
            style.font.bold = bold
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

    set_style_font("Normal", body_font, body_size, False)
    set_style_font("Heading 1", heading_font, _pt_value(_style_value(style_profile, "heading_1_size", "16pt"), 16), True)
    set_style_font("Heading 2", heading_font, _pt_value(_style_value(style_profile, "heading_2_size", "14pt"), 14), True)
    set_style_font("Heading 3", heading_font, _pt_value(_style_value(style_profile, "heading_3_size", "12pt"), 12), True)


@router.post("/upload", response_model=FileUploadResponse)
async def upload_file(file: UploadFile = File(...)):
    """上传文档文件并提取文本内容"""
    try:
        file_kind = FileService.detect_upload_file_kind(file)
        if file_kind in (None, "doc"):
            return FileUploadResponse(
                success=False,
                message=FileService.get_upload_validation_message(file)
            )
        
        # 处理文件并提取文本
        parse_result = await FileService.process_uploaded_file_with_metadata(file)
        file_content = parse_result.get("file_content", "")
        parser_info = parse_result.get("parser_info", {})
        parser_name = parser_info.get("parser") or "unknown"
        fallback_note = "（已降级到内置解析器）" if parser_info.get("fallback_used") else ""
        
        return FileUploadResponse(
            success=True,
            message=f"文件 {file.filename} 上传成功，解析器：{parser_name}{fallback_note}，已在上传阶段转换为 {parser_info.get('format') or '文本'}",
            file_content=file_content,
            parser_info=parser_info,
        )
        
    except HTTPException:
        raise
    except Exception as e:
        return FileUploadResponse(
            success=False,
            message=f"文件处理失败: {str(e)}"
        )


@router.post("/analyze-stream")
async def analyze_document_stream(request: AnalysisRequest):
    """流式分析文档内容"""
    try:
        # 加载配置
        config = config_manager.load_config()
        auth_error = get_provider_auth_error(config.get("provider"), config.get("api_key"))
        if auth_error:
            raise HTTPException(status_code=400, detail=auth_error)

        # 创建OpenAI服务实例
        openai_service = OpenAIService()
        
        async def generate():
            try:
                if openai_service._force_local_fallback():
                    fallback_text = (
                        OpenAIService.fallback_overview(request.file_content)
                        if request.analysis_type == AnalysisType.OVERVIEW
                        else OpenAIService.fallback_requirements(request.file_content)
                    )
                    yield f"data: {json.dumps({'chunk': fallback_text}, ensure_ascii=False)}\n\n"
                    yield "data: [DONE]\n\n"
                    return

                # 构建分析提示词
                if request.analysis_type == AnalysisType.OVERVIEW:
                    system_prompt = """你是一个专业的标书撰写专家。请分析用户发来的招标文件，提取并总结项目概述信息。
            
请重点关注以下方面：
1. 项目名称和基本信息
2. 项目背景和目的
3. 项目规模和预算
4. 项目时间安排
5. 项目要实施的具体内容
6. 主要技术特点
7. 其他关键要求

工作要求：
1. 保持提取信息的全面性和准确性，尽量使用原文内容，不要自己编写
2. 只关注与项目实施有关的内容，不提取商务信息
3. 直接返回整理好的项目概述，除此之外不返回任何其他内容
"""
                else:  # requirements
                    system_prompt = """你是一名专业的招标文件分析师，擅长从复杂的招标文档中高效提取“技术评分项”相关内容。请严格按照以下步骤和规则执行任务：
### 1. 目标定位
- 重点识别文档中与“技术评分”、“评标方法”、“评分标准”、“技术参数”、“技术要求”、“技术方案”、“技术部分”或“评审要素”相关的章节（如“第X章 评标方法”或“附件X：技术评分表”）。
- 一定不要提取商务、价格、资质等于技术类评分项无关的条目。
### 2. 提取内容要求
对每一项技术评分项，按以下结构化格式输出（若信息缺失，标注“未提及”），如果评分项不够明确，你需要根据上下文分析并也整理成如下格式：
【评分项名称】：<原文描述，保留专业术语>
【权重/分值】：<具体分值或占比，如“30分”或“40%”>
【评分标准】：<详细规则，如“≥95%得满分，每低1%扣0.5分”>
【数据来源】：<文档中的位置，如“第5.2.3条”或“附件3-表2”>

### 3. 处理规则
- **模糊表述**：有些招标文件格式不是很标准，没有明确的“技术评分表”，但一定都会有“技术评分”相关内容，请根据上下文判断评分项。
- **表格处理**：若评分项以表格形式呈现，按行提取，并标注“[表格数据]”。
- **分层结构**：若存在二级评分项（如“技术方案→子项1、子项2”），用缩进或编号体现层级关系。
- **单位统一**：将所有分值统一为“分”或“%”，并注明原文单位（如原文为“20点”则标注“[原文：20点]”）。

### 4. 输出示例
【评分项名称】：系统可用性 
【权重/分值】：25分 
【评分标准】：年平均故障时间≤1小时得满分；每增加1小时扣2分，最高扣10分。 
【数据来源】：附件4-技术评分细则（第3页） 

【评分项名称】：响应时间
【权重/分分】：15分 [原文：15%]
【评分标准】：≤50ms得满分；每增加10ms扣1分。
【数据来源】：第6.1.2条

### 5. 验证步骤
提取完成后，执行以下自检：
- [ ] 所有技术评分项是否覆盖（无遗漏）？
- [ ] 是否错误提取商务、价格、资质等于技术类评分项无关的条目？
- [ ] 权重总和是否与文档声明的技术分总分一致（如“技术部分共60分”）？

直接返回提取结果，除此之外不输出任何其他内容
"""
                
                analysis_type_cn = "项目概述" if request.analysis_type == AnalysisType.OVERVIEW else "技术评分要求"
                user_prompt = f"请分析以下招标文件内容，提取{analysis_type_cn}信息：\n\n{request.file_content}"
                
                messages = [
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ]
                
                # 流式返回分析结果
                try:
                    async for chunk in openai_service.stream_chat_completion(messages, temperature=0.3):
                        yield f"data: {json.dumps({'chunk': chunk}, ensure_ascii=False)}\n\n"
                except Exception as e:
                    fallback_text = (
                        OpenAIService.fallback_overview(request.file_content)
                        if request.analysis_type == AnalysisType.OVERVIEW
                        else OpenAIService.fallback_requirements(request.file_content)
                    )
                    print(f"文档分析模型输出不可用，启用文本兜底分析：{str(e)}")
                    yield f"data: {json.dumps({'chunk': fallback_text}, ensure_ascii=False)}\n\n"
            except Exception as e:
                payload = {
                    "error": True,
                    "message": f"文档分析失败: {str(e)}",
                }
                yield f"data: {json.dumps(payload, ensure_ascii=False)}\n\n"

            # 发送结束信号
            yield "data: [DONE]\n\n"
        
        return sse_response(generate())
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"文档分析失败: {str(e)}")


@router.post("/analyze-report-stream")
async def analyze_report_stream(request: AnalysisReportRequest):
    """流式生成结构化标准解析报告"""
    try:
        request_id = uuid.uuid4().hex[:8]
        started_at = time.monotonic()
        task_state = AnalysisTaskState(request_id)
        ANALYSIS_TASKS[request_id] = task_state
        config = (
            request.config.model_dump(mode="json", exclude_none=True)
            if request.config
            else config_manager.load_config()
        )
        auth_error = get_provider_auth_error(config.get("provider"), config.get("api_key"))
        if auth_error:
            raise HTTPException(status_code=400, detail=auth_error)

        openai_service = OpenAIService(config=config)
        print(
            f"标准解析报告[{request_id}] 开始："
            f"model={config.get('model_name')} base_url={config.get('base_url')} "
            f"file_chars={len(request.file_content or '')}",
            flush=True,
        )

        async def generate():
            compute_task = None
            final_report = None

            async def wait_if_paused() -> None:
                while not task_state.pause_event.is_set():
                    task_state.status = "paused"
                    if task_state.cancelled:
                        raise asyncio.CancelledError()
                    await asyncio.sleep(0.25)
                if task_state.cancelled:
                    raise asyncio.CancelledError()
                task_state.status = "running"

            try:
                yield _analysis_stream_payload(
                    task_state,
                    0,
                    "文件解析完成：MinerU Markdown 已进入标准解析队列",
                    12,
                )
                await wait_if_paused()
                yield _analysis_stream_payload(
                    task_state,
                    1,
                    "条款识别中：正在抽取项目基础信息和投标文件组成",
                    20,
                )
                compute_task = asyncio.create_task(
                    openai_service.generate_analysis_report(request.file_content)
                )
                heartbeat_count = 0

                while not compute_task.done():
                    await wait_if_paused()
                    yield _analysis_stream_payload(
                        task_state,
                        1,
                        "条款识别中：正在抽取项目基础信息和投标文件组成",
                        20,
                    )
                    heartbeat_count += 1
                    if heartbeat_count % 30 == 0:
                        print(
                            f"标准解析报告[{request_id}] 仍在生成："
                            f"{int(time.monotonic() - started_at)} 秒，心跳 {heartbeat_count} 次",
                            flush=True,
                        )
                    await asyncio.sleep(1)

                await wait_if_paused()
                final_report = await compute_task
                yield _analysis_stream_payload(
                    task_state,
                    1,
                    "条款识别完成：项目基础信息、投标文件组成和关键出处已抽取",
                    38,
                )
                await wait_if_paused()

                scoring_count = _analysis_metric_count(
                    final_report,
                    ("technical_scoring_items", "business_scoring_items", "price_scoring_items"),
                )
                yield _analysis_stream_payload(
                    task_state,
                    2,
                    f"评分项提取完成：识别 {scoring_count} 项技术、商务或报价评分规则",
                    58,
                )
                await wait_if_paused()

                compliance_count = _analysis_metric_count(
                    final_report,
                    (
                        "qualification_requirements",
                        "formal_response_requirements",
                        "mandatory_clauses",
                        "rejection_risks",
                        "required_materials",
                        "fixed_format_forms",
                        "signature_requirements",
                    ),
                )
                yield _analysis_stream_payload(
                    task_state,
                    3,
                    f"合规要求提取完成：识别 {compliance_count} 项资格、格式、签章、材料或废标风险",
                    78,
                )
                await wait_if_paused()

                yield _analysis_stream_payload(
                    task_state,
                    4,
                    "结果校验中：正在校验 JSON、响应矩阵和企业资料画像",
                    92,
                )
                report_json = json.dumps(final_report, ensure_ascii=False)
                print(
                    f"标准解析报告[{request_id}] 完成："
                    f"{int(time.monotonic() - started_at)} 秒，输出 {len(report_json)} 字符",
                    flush=True,
                )
                chunk_size = 256
                for index in range(0, len(report_json), chunk_size):
                    await wait_if_paused()
                    piece = report_json[index:index + chunk_size]
                    yield f"data: {json.dumps({'chunk': piece, 'task_id': request_id, 'step_index': 4, 'percent': 96, 'status': 'running'}, ensure_ascii=False)}\n\n"
                yield _analysis_stream_payload(
                    task_state,
                    4,
                    "标准解析完成：已写入项目、评分、风险和材料结构",
                    100,
                    "success",
                )
            except asyncio.CancelledError:
                if compute_task and not compute_task.done():
                    compute_task.cancel()
                if task_state.cancelled:
                    print(f"标准解析报告[{request_id}] 已停止", flush=True)
                    yield _analysis_stream_payload(
                        task_state,
                        task_state.current_step,
                        "标准解析已停止，可重新开始解析",
                        max(0, min(100, 0)),
                        "stopped",
                        stopped=True,
                    )
                else:
                    print(f"标准解析报告[{request_id}] SSE 连接被取消", flush=True)
                    raise
            except Exception as e:
                print(f"标准解析报告[{request_id}] 失败：{str(e)}", flush=True)
                payload = {
                    "chunk": "",
                    "task_id": request_id,
                    "step_index": task_state.current_step,
                    "percent": 100,
                    "status": "error",
                    "error": True,
                    "message": f"结构化解析报告生成失败: {str(e)}",
                }
                yield f"data: {json.dumps(payload, ensure_ascii=False)}\n\n"
            finally:
                ANALYSIS_TASKS.pop(request_id, None)

            yield "data: [DONE]\n\n"

        return sse_response(generate())
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"结构化解析报告生成失败: {str(e)}")


@router.post("/analysis-task/{task_id}/control", response_model=AnalysisTaskControlResponse)
async def control_analysis_task(task_id: str, request: AnalysisTaskControlRequest):
    """Pause, resume, or stop a running standard-analysis task."""
    task_state = ANALYSIS_TASKS.get(task_id)
    if not task_state:
        raise HTTPException(status_code=404, detail="解析任务不存在或已结束")

    action = (request.action or "").strip().lower()
    if action == "pause":
        task_state.pause_event.clear()
        task_state.status = "paused"
        return AnalysisTaskControlResponse(success=True, message="标准解析已暂停", task_id=task_id, status=task_state.status)
    if action in {"resume", "continue"}:
        task_state.pause_event.set()
        task_state.status = "running"
        return AnalysisTaskControlResponse(success=True, message="标准解析已继续", task_id=task_id, status=task_state.status)
    if action == "stop":
        task_state.cancelled = True
        task_state.pause_event.set()
        task_state.status = "stopped"
        return AnalysisTaskControlResponse(success=True, message="标准解析已停止", task_id=task_id, status=task_state.status)

    raise HTTPException(status_code=400, detail="不支持的解析任务控制动作")


@router.post("/reference-style-upload")
async def reference_style_upload(file: UploadFile = File(...)):
    """上传成熟投标文件样例并生成可复用风格剖面"""
    try:
        file_kind = FileService.detect_upload_file_kind(file)
        if file_kind in (None, "doc"):
            return FileUploadResponse(success=False, message=FileService.get_upload_validation_message(file))

        parse_result = await FileService.process_uploaded_file_with_metadata(file)
        file_content = parse_result.get("file_content", "")
        parser_info = parse_result.get("parser_info", {})
        if (
            str(parser_info.get("parser") or "").lower() != "mineru"
            or str(parser_info.get("format") or "").lower() != "markdown"
        ):
            return FileUploadResponse(
                success=False,
                message="成熟样例必须先通过 MinerU 转换为 Markdown 后再生成模板剖面；请检查 MinerU 配置或将 YIBIAO_DOCUMENT_PARSER 设置为 mineru_strict。",
                parser_info=parser_info,
            )

        config = config_manager.load_config()
        auth_error = get_provider_auth_error(config.get("provider"), config.get("api_key"))
        if auth_error:
            raise HTTPException(status_code=400, detail=auth_error)

        profile = await OpenAIService().generate_reference_bid_style_profile(file_content)
        return FileUploadResponse(
            success=True,
            message=f"样例文件 {file.filename} 已通过 MinerU Markdown 解析，并生成写作模板剖面",
            file_content=file_content,
            parser_info=parser_info,
            reference_bid_style_profile=profile,
        )
    except HTTPException as e:
        return FileUploadResponse(success=False, message=e.detail)
    except Exception as e:
        return FileUploadResponse(success=False, message=f"样例解析失败: {str(e)}")


@router.post("/document-blocks-plan-stream")
async def document_blocks_plan_stream(request: DocumentBlocksPlanRequest):
    """生成图表、表格、图片、承诺书和附件规划"""
    try:
        config = config_manager.load_config()
        auth_error = get_provider_auth_error(config.get("provider"), config.get("api_key"))
        if auth_error:
            raise HTTPException(status_code=400, detail=auth_error)

        async def generate():
            try:
                service = OpenAIService()
                compute_task = asyncio.create_task(service.generate_document_blocks_plan(
                    outline=[item.model_dump(mode="json") for item in request.outline],
                    analysis_report=request.analysis_report.model_dump(mode="json") if request.analysis_report else None,
                    response_matrix=request.response_matrix.model_dump(mode="json") if request.response_matrix else None,
                    reference_bid_style_profile=request.reference_bid_style_profile,
                    enterprise_materials=[item.model_dump(mode="json") for item in request.enterprise_materials],
                    asset_library=request.asset_library,
                ))
                while not compute_task.done():
                    yield f"data: {json.dumps({'chunk': ''}, ensure_ascii=False)}\n\n"
                    await asyncio.sleep(1)
                result_json = json.dumps(await compute_task, ensure_ascii=False)
                for index in range(0, len(result_json), 256):
                    yield f"data: {json.dumps({'chunk': result_json[index:index + 256]}, ensure_ascii=False)}\n\n"
            except Exception as e:
                yield f"data: {json.dumps({'chunk': '', 'error': True, 'message': f'图表素材规划失败: {str(e)}'}, ensure_ascii=False)}\n\n"
            yield "data: [DONE]\n\n"

        return sse_response(generate())
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"图表素材规划失败: {str(e)}")


@router.post("/consistency-revision-stream")
async def consistency_revision_stream(request: ConsistencyRevisionRequest):
    """生成全文一致性修订报告"""
    try:
        config = config_manager.load_config()
        auth_error = get_provider_auth_error(config.get("provider"), config.get("api_key"))
        if auth_error:
            raise HTTPException(status_code=400, detail=auth_error)

        async def generate():
            try:
                service = OpenAIService()
                compute_task = asyncio.create_task(service.generate_consistency_revision_report(
                    full_bid_draft=[item.model_dump(mode="json") for item in request.full_bid_draft],
                    analysis_report=request.analysis_report.model_dump(mode="json") if request.analysis_report else None,
                    response_matrix=request.response_matrix.model_dump(mode="json") if request.response_matrix else None,
                    reference_bid_style_profile=request.reference_bid_style_profile,
                    document_blocks_plan=request.document_blocks_plan,
                ))
                while not compute_task.done():
                    yield f"data: {json.dumps({'chunk': ''}, ensure_ascii=False)}\n\n"
                    await asyncio.sleep(1)
                result_json = json.dumps(await compute_task, ensure_ascii=False)
                for index in range(0, len(result_json), 256):
                    yield f"data: {json.dumps({'chunk': result_json[index:index + 256]}, ensure_ascii=False)}\n\n"
            except Exception as e:
                yield f"data: {json.dumps({'chunk': '', 'error': True, 'message': f'全文一致性修订失败: {str(e)}'}, ensure_ascii=False)}\n\n"
            yield "data: [DONE]\n\n"

        return sse_response(generate())
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"全文一致性修订失败: {str(e)}")


@router.post("/review-compliance-stream")
async def review_compliance_stream(request: ComplianceReviewRequest):
    """导出前执行覆盖性、缺料、废标风险和虚构风险审校"""
    try:
        config = config_manager.load_config()
        auth_error = get_provider_auth_error(config.get("provider"), config.get("api_key"))
        if auth_error:
            raise HTTPException(status_code=400, detail=auth_error)

        openai_service = OpenAIService()

        async def generate():
            try:
                compute_task = asyncio.create_task(
                    openai_service.generate_compliance_review(
                        outline=[item.model_dump(mode="json") for item in request.outline],
                        analysis_report=(
                            request.analysis_report.model_dump(mode="json")
                            if request.analysis_report else None
                        ),
                        response_matrix=(
                            request.response_matrix.model_dump(mode="json")
                            if request.response_matrix else None
                        ),
                        project_overview=request.project_overview or "",
                        reference_bid_style_profile=request.reference_bid_style_profile,
                        document_blocks_plan=request.document_blocks_plan,
                    )
                )

                while not compute_task.done():
                    yield f"data: {json.dumps({'chunk': ''}, ensure_ascii=False)}\n\n"
                    await asyncio.sleep(1)

                review_json = json.dumps(await compute_task, ensure_ascii=False)
                chunk_size = 256
                for index in range(0, len(review_json), chunk_size):
                    piece = review_json[index:index + chunk_size]
                    yield f"data: {json.dumps({'chunk': piece}, ensure_ascii=False)}\n\n"
            except Exception as e:
                payload = {
                    "chunk": "",
                    "error": True,
                    "message": f"导出前合规审校失败: {str(e)}",
                }
                yield f"data: {json.dumps(payload, ensure_ascii=False)}\n\n"

            yield "data: [DONE]\n\n"

        return sse_response(generate())
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"导出前合规审校失败: {str(e)}")


@router.post("/export-word")
async def export_word(request: WordExportRequest):
    """根据目录数据导出Word文档"""
    try:
        if not request.manual_review_confirmed:
            raise HTTPException(status_code=400, detail="导出前必须完成人工复核确认，不能直接跳过模型结果复核。")

        doc = docx.Document()
        enable_update_fields_on_open(doc)
        configure_word_style(doc, request.reference_bid_style_profile)

        # 统一设置普通正文基础字体，标题样式已由成熟样例模板配置。
        try:
            styles = doc.styles
            base_styles = ["Normal"]
            for style_name in base_styles:
                if style_name in styles:
                    style = styles[style_name]
                    font = style.font
                    font.name = "宋体"
                    # 设置中文字体
                    if style._element.rPr is None:
                        style._element._add_rPr()
                    rpr = style._element.rPr
                    rpr.rFonts.set(qn("w:eastAsia"), "宋体")
                    if style_name == "Normal":
                        font.bold = False
        except Exception:
            # 字体设置失败不影响文档生成，忽略
            pass

        # AI 生成声明
        p = doc.add_paragraph()
        run = p.add_run("内容由 AI 生成，导出前已要求人工复核确认；最终页码、目录、签章、版式和图表需在 Word 中复核。")
        run.italic = True
        run.font.size = Pt(9)
        set_run_font_simsun(run)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        footer_p = doc.sections[0].footer.paragraphs[0]
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_p.add_run("第 ")
        add_field(footer_p, "PAGE")
        footer_p.add_run(" 页 / 共 ")
        add_field(footer_p, "NUMPAGES")
        footer_p.add_run(" 页")
        set_paragraph_font_simsun(footer_p)

        # 文档标题
        title = request.project_name or "投标技术文件"
        title_p = doc.add_paragraph()
        title_run = title_p.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(16)
        set_run_font_simsun(title_run)
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 可更新目录。python-docx 不能计算最终页码，但可以生成 Word 域。
        toc_heading = doc.add_heading("目录", level=1)
        set_paragraph_font_simsun(toc_heading)
        toc_p = doc.add_paragraph()
        add_field(toc_p, r'TOC \o "1-3" \h \z \u')
        doc.add_page_break()

        def add_export_checklist() -> None:
            """Add a visible finalization checklist for manual Word work."""
            report = request.analysis_report.model_dump(mode="json") if request.analysis_report else {}
            review = request.review_report.model_dump(mode="json") if request.review_report else {}
            bid_rules = report.get("bid_document_requirements") or {}
            formatting = bid_rules.get("formatting_and_submission_rules") or {}
            signature_items = report.get("signature_requirements") or []
            fixed_forms = report.get("fixed_format_forms") or []
            enterprise_profile = report.get("enterprise_material_profile") or {}
            missing_enterprise = enterprise_profile.get("missing_materials") or report.get("missing_company_materials") or []
            plan = request.document_blocks_plan or {}
            missing_assets = []
            if isinstance(plan, dict):
                missing_assets = (plan.get("missing_assets") or []) + (plan.get("missing_enterprise_data") or [])
            blocking_count = ((review.get("summary") or {}).get("blocking_issues_count")
                              or (review.get("summary") or {}).get("blocking_issues") or 0)

            heading = doc.add_heading("导出后 Word 处理清单", level=1)
            set_paragraph_font_simsun(heading)
            lines = [
                "更新目录域和页码域：打开 Word 后全选并更新域，确认目录层级和页码正确。",
                "复核版式：按招标文件检查页边距、字体、行距、标题层级、表格跨页和页眉页脚。",
                "复核签章位置：按签章要求在投标函、授权委托书、报价表、承诺函、固定格式表单等位置签字盖章。",
                "复核固定格式：不得破坏招标文件固定表头、固定文字、行列结构和附件说明。",
                "复核图表与附件：将图表、组织架构图、流程图、证书、截图、扫描件替换到对应占位。",
                "复核模型风险：逐项核对误读、漏项、虚构、历史项目残留和格式偏差；不能仅凭模型审校结论提交。",
            ]
            if formatting.get("toc_required") or formatting.get("page_number_required"):
                lines.append("招标文件要求目录/页码，请在最终版中更新目录和所有响应页码。")
            if signature_items:
                lines.append(f"已解析签章要求 {len(signature_items)} 项，需逐项核验签署主体、盖章、日期和电子签章。")
            if fixed_forms:
                lines.append(f"已解析固定格式 {len(fixed_forms)} 项，需人工确认表头、固定文字和签章栏。")
            if missing_enterprise:
                lines.append(f"企业资料仍有 {len(missing_enterprise)} 项待补或待确认，正文中的待补占位不得直接提交。")
            if missing_assets:
                lines.append(f"图表/素材/企业数据仍有 {len(missing_assets)} 项待替换。")
            if blocking_count:
                lines.append(f"审校报告仍有 {blocking_count} 项阻塞问题，处理后再提交最终版。")

            for line in lines:
                item_p = doc.add_paragraph()
                run = item_p.add_run(f"□ {line}")
                set_run_font_simsun(run)
            doc.add_page_break()

        add_export_checklist()

        # 项目概述
        if request.project_overview:
            heading = doc.add_heading("项目概述", level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_font_simsun(heading)
            overview_p = doc.add_paragraph(request.project_overview)
            set_paragraph_font_simsun(overview_p)
            overview_p_format = overview_p.paragraph_format
            overview_p_format.space_after = Pt(12)

        # 简单的 Markdown 段落解析：支持标题、列表、表格和基础加粗/斜体
        def add_markdown_runs(para: docx.text.paragraph.Paragraph, text: str) -> None:
            """在指定段落中追加 markdown 文本的 runs"""
            pattern = r"(\*\*.*?\*\*|\*.*?\*|`.*?`)"
            parts = re.split(pattern, text)
            for part in parts:
                if not part:
                    continue
                run = para.add_run()
                # 加粗
                if part.startswith("**") and part.endswith("**") and len(part) > 4:
                    run.text = part[2:-2]
                    run.bold = True
                # 斜体
                elif part.startswith("*") and part.endswith("*") and len(part) > 2:
                    run.text = part[1:-1]
                    run.italic = True
                # 行内代码：这里只去掉反引号
                elif part.startswith("`") and part.endswith("`") and len(part) > 2:
                    run.text = part[1:-1]
                else:
                    run.text = part
                # 确保字体为宋体
                set_run_font_simsun(run)

        def add_markdown_paragraph(text: str) -> None:
            """将一段 Markdown 文本解析为一个普通段落，保留加粗/斜体效果"""
            para = doc.add_paragraph()
            add_markdown_runs(para, text)
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.first_line_indent = Pt(21)
            para.paragraph_format.line_spacing = 1.5

        def parse_markdown_blocks(content: str):
            """
            识别 Markdown 内容中的块级元素，返回结构化的 block 列表：
            - ('list', items)        items: [(kind, num_str, text), ...]
            - ('table', rows)        rows: [text, ...]
            - ('heading', level, text)
            - ('paragraph', text)
            """
            blocks = []
            lines = content.split("\n")
            i = 0
            while i < len(lines):
                line = lines[i].rstrip("\r").strip()
                if not line:
                    i += 1
                    continue

                # 列表项（有序/无序）
                if line.startswith("- ") or line.startswith("* ") or re.match(r"^\d+\.\s", line):
                    # items: (kind, number, text)
                    items = []
                    while i < len(lines):
                        raw = lines[i].rstrip("\r")
                        stripped = raw.strip()
                        # 无序列表
                        if stripped.startswith("- ") or stripped.startswith("* "):
                            text = re.sub(r"^[-*]\s+", "", stripped).strip()
                            if text:
                                items.append(("unordered", None, text))
                            i += 1
                            continue
                        # 有序列表（1. xxx）
                        m_num = re.match(r"^(\d+)\.\s+(.*)$", stripped)
                        if m_num:
                            num_str, text = m_num.groups()
                            text = text.strip()
                            if text:
                                items.append(("ordered", num_str, text))
                            i += 1
                            continue
                        break

                    if items:
                        blocks.append(("list", items))
                    continue

                # 表格（简化为每行一个段落，单元格用 | 分隔）
                if "|" in line:
                    rows = []
                    while i < len(lines):
                        raw = lines[i].rstrip("\r")
                        stripped = raw.strip()
                        if "|" in stripped:
                            # 跳过仅由 - 和 | 组成的分隔行
                            if not re.match(r"^\|?[-\s\|]+\|?$", stripped):
                                cells = [c.strip() for c in stripped.split("|")]
                                row_text = " | ".join([c for c in cells if c])
                                if row_text:
                                    rows.append(row_text)
                            i += 1
                        else:
                            break
                    if rows:
                        blocks.append(("table", rows))
                    continue

                # Markdown 标题（# / ## / ###）
                if line.startswith("#"):
                    m = re.match(r"^(#+)\s*(.*)$", line)
                    if m:
                        level_marks, title_text = m.groups()
                        level = min(len(level_marks), 3)
                        blocks.append(("heading", level, title_text.strip()))
                    i += 1
                    continue

                # 普通段落：合并连续的普通行
                para_lines = []
                while i < len(lines):
                    raw = lines[i].rstrip("\r")
                    stripped = raw.strip()
                    if (
                        stripped
                        and not stripped.startswith("-")
                        and not stripped.startswith("*")
                        and "|" not in stripped
                        and not stripped.startswith("#")
                    ):
                        para_lines.append(stripped)
                        i += 1
                    else:
                        break
                if para_lines:
                    text = " ".join(para_lines)
                    blocks.append(("paragraph", text))
                else:
                    i += 1

            return blocks

        def render_markdown_blocks(blocks) -> None:
            """将结构化的 Markdown blocks 渲染到文档"""
            for block in blocks:
                kind = block[0]
                if kind == "list":
                    items = block[1]
                    for item_kind, num_str, text in items:
                        p = doc.add_paragraph()
                        if item_kind == "unordered":
                            # 使用“• ”模拟项目符号
                            run = p.add_run("• ")
                            set_run_font_simsun(run)
                        else:
                            # 有序列表：输出 "1. " 这样的前缀
                            prefix = f"{num_str}."
                            run = p.add_run(prefix + " ")
                            set_run_font_simsun(run)
                        # 紧跟在同一段落中追加列表文本
                        add_markdown_runs(p, text)
                        p.paragraph_format.space_after = Pt(6)
                        p.paragraph_format.line_spacing = 1.5
                elif kind == "table":
                    rows = block[1]
                    for row in rows:
                        add_markdown_paragraph(row)
                elif kind == "heading":
                    _, level, text = block
                    heading = doc.add_heading(text, level=level)
                    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    set_paragraph_font_simsun(heading)
                elif kind == "paragraph":
                    _, text = block
                    add_markdown_paragraph(text)

        def add_markdown_content(content: str) -> None:
            """解析并渲染 Markdown 文本到文档"""
            blocks = parse_markdown_blocks(content)
            render_markdown_blocks(blocks)

        def planned_blocks_by_chapter() -> dict[str, list[dict]]:
            """按章节聚合图表/素材规划，导出时渲染为可替换占位。"""
            plan = request.document_blocks_plan or {}
            groups: dict[str, list[dict]] = {}
            raw_blocks = plan.get("document_blocks") if isinstance(plan, dict) else []
            for block in raw_blocks or []:
                if not isinstance(block, dict):
                    continue
                chapter_id = str(block.get("chapter_id") or block.get("target_chapter_id") or "").strip()
                if not chapter_id:
                    continue
                groups.setdefault(chapter_id, []).append(block)
            return groups

        blocks_by_chapter = planned_blocks_by_chapter()

        def add_planned_blocks(chapter_id: str) -> None:
            """将文档块规划渲染为 Word 中的表格/图片/承诺书占位。"""
            planned = blocks_by_chapter.get(str(chapter_id), [])
            if not planned:
                return
            heading = doc.add_heading("文档块规划", level=4)
            set_paragraph_font_simsun(heading)
            for block in planned:
                block_name = block.get("block_name") or block.get("name") or block.get("asset_name") or "文档块"
                block_type = block.get("block_type") or block.get("type") or "block"
                placeholder = block.get("placeholder") or block.get("fallback_placeholder") or ""
                p = doc.add_paragraph()
                add_markdown_runs(p, f"【{block_type}】{block_name}")
                if placeholder:
                    add_markdown_paragraph(str(placeholder))

                table_schema = block.get("table_schema") or {}
                columns = table_schema.get("columns") if isinstance(table_schema, dict) else []
                if columns:
                    table = doc.add_table(rows=2, cols=len(columns))
                    table.style = "Table Grid"
                    for idx, column in enumerate(columns):
                        table.rows[0].cells[idx].text = str(column)
                        table.rows[1].cells[idx].text = "〖待补充〗"
                        for paragraph in table.rows[0].cells[idx].paragraphs + table.rows[1].cells[idx].paragraphs:
                            set_paragraph_font_simsun(paragraph)

                chart_schema = block.get("chart_schema") or {}
                if isinstance(chart_schema, dict) and (chart_schema.get("nodes") or chart_schema.get("edges")):
                    add_markdown_paragraph(f"〖插入图表：{block_name}〗")

                commitment_schema = block.get("commitment_schema") or {}
                if isinstance(commitment_schema, dict) and commitment_schema.get("items"):
                    add_markdown_paragraph(f"承诺书：{block_name}")
                    for item in commitment_schema.get("items", []):
                        add_markdown_paragraph(f"- {item}")

        # 递归构建文档内容（章节和内容）
        def add_outline_items(items, level: int = 1):
            for item in items:
                # 章节标题
                if level <= 3:
                    heading = doc.add_heading(f"{item.id} {item.title}", level=level)
                    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for hr in heading.runs:
                        hr.font.name = "宋体"
                        rr = hr._element.rPr
                        if rr is not None and rr.rFonts is not None:
                            rr.rFonts.set(qn("w:eastAsia"), "宋体")
                else:
                    para = doc.add_paragraph()
                    run = para.add_run(f"{item.id} {item.title}")
                    run.bold = True
                    run.font.name = "宋体"
                    rr = run._element.rPr
                    if rr is not None and rr.rFonts is not None:
                        rr.rFonts.set(qn("w:eastAsia"), "宋体")
                    para.paragraph_format.space_before = Pt(6)
                    para.paragraph_format.space_after = Pt(3)

                # 叶子节点内容
                if not item.children:
                    content = item.content or ""
                    if content.strip():
                        add_markdown_content(content)
                    add_planned_blocks(item.id)
                else:
                    add_outline_items(item.children, level + 1)

        add_outline_items(request.outline)

        # 输出到内存并返回
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        filename = f"{request.project_name or '标书文档'}.docx"
        if request.export_dir:
            export_dir = Path(request.export_dir).expanduser()
            export_dir.mkdir(parents=True, exist_ok=True)
            if not export_dir.is_dir():
                raise RuntimeError(f"保存路径不是目录：{export_dir}")

            safe_filename = sanitize_docx_filename(filename)
            output_path = export_dir / safe_filename
            output_path.write_bytes(buffer.getvalue())
            return JSONResponse({
                "success": True,
                "message": "Word 文件已保存到指定目录",
                "file_path": str(output_path),
                "filename": safe_filename,
            })

        # 使用 RFC 5987 格式对文件名进行 URL 编码，避免非 ASCII 字符导致的编码错误
        encoded_filename = quote(filename)
        content_disposition = f"attachment; filename*=UTF-8''{encoded_filename}"
        headers = {
            "Content-Disposition": content_disposition
        }

        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers=headers
        )
    except HTTPException:
        raise
    except Exception as e:
        # 打印详细错误信息到控制台，方便排查
        import traceback
        print("导出Word失败:", str(e))
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"导出Word失败: {str(e)}")
