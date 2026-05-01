"""文件处理服务"""
import aiofiles
import json
import os
import platform
import shutil
import subprocess
import time
import gc
import io
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Optional, List, Dict, Tuple, Any
import PyPDF2
import docx
from fastapi import UploadFile
import aiohttp
import asyncio
from ..config import settings

# 新增的第三方库
try:
    import pdfplumber
    import fitz  # PyMuPDF
    from docx2python import docx2python
    HAS_ADVANCED_LIBS = True
except ImportError as e:
    HAS_ADVANCED_LIBS = False
    print(f"高级文档处理库未安装: {e}")


class FileService:
    """文件处理服务"""

    SUPPORTED_UPLOAD_MESSAGE = "仅支持 PDF 和 DOCX 文件，暂不支持 DOC，请先另存为 DOCX 后再上传"
    MINERU_DEFAULT_TIMEOUT = 900

    # 图片上传配置
    IMAGE_UPLOAD_URL = "https://mt.agnet.top/image/upload"
    IMAGE_UPLOAD_TIMEOUT = 30  # 超时时间（秒）

    @staticmethod
    def get_parser_mode() -> str:
        """文档解析器选择：auto / mineru / mineru_strict / legacy。"""
        mode = os.getenv("YIBIAO_DOCUMENT_PARSER", "auto").strip().lower()
        if mode not in {"auto", "mineru", "mineru_strict", "legacy"}:
            return "auto"
        return mode

    @staticmethod
    def _mineru_command() -> Optional[str]:
        """返回本机可用 MinerU CLI 路径。"""
        configured = os.getenv("YIBIAO_MINERU_BIN", "mineru").strip() or "mineru"
        if os.path.isabs(configured) and os.access(configured, os.X_OK):
            return configured
        resolved = shutil.which(configured)
        if resolved:
            return resolved

        mineru_home = os.getenv("YIBIAO_MINERU_HOME", "").strip()
        candidate_roots = [Path(mineru_home)] if mineru_home else []
        current_path = Path(__file__).resolve()
        candidate_roots.extend([
            current_path.parents[6] / "MinerU",
            current_path.parents[5] / "MinerU",
            Path("/Users/songyuheng/Documents/01work/python_code_locate/MinerU"),
        ])
        for root in candidate_roots:
            for relative in (".venv-mps/bin/mineru", ".venv/bin/mineru"):
                candidate = root / relative
                if candidate.exists() and os.access(candidate, os.X_OK):
                    return str(candidate)
        return None

    @staticmethod
    def _detect_mineru_device() -> str:
        """选择 MinerU 本地推理设备：CUDA -> MPS -> CPU。"""
        configured = os.getenv("YIBIAO_MINERU_DEVICE", "auto").strip().lower()
        if configured in {"cuda", "mps", "cpu"}:
            return configured

        if shutil.which("nvidia-smi"):
            try:
                proc = subprocess.run(["nvidia-smi", "-L"], timeout=2, capture_output=True, text=True)
                if proc.returncode == 0 and proc.stdout.strip():
                    return "cuda"
            except Exception:
                pass

        if platform.system().lower() == "darwin":
            return "mps"
        return "cpu"

    @staticmethod
    def detect_upload_file_kind(file: UploadFile) -> Optional[str]:
        """根据扩展名和 MIME 类型识别上传文件类型"""
        filename = (file.filename or "").lower()
        extension = os.path.splitext(filename)[1]
        content_type = (file.content_type or "").lower()

        if extension == ".doc" or content_type in {
            "application/msword",
            "application/doc",
            "application/vnd.ms-word",
        }:
            return "doc"

        if extension == ".pdf" or content_type == "application/pdf":
            return "pdf"

        if extension == ".docx" or content_type == (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ):
            return "docx"

        return None

    @staticmethod
    def get_upload_validation_message(file: UploadFile) -> str:
        """返回上传文件类型校验失败时的提示信息"""
        if FileService.detect_upload_file_kind(file) == "doc":
            return "暂不支持 .doc 文件，请先另存为 .docx 后再上传"
        return FileService.SUPPORTED_UPLOAD_MESSAGE

    @staticmethod
    async def upload_image_to_server(image_data: bytes, filename: str) -> Optional[str]:
        """上传图片到外部服务器"""
        try:
            # 准备multipart/form-data格式的数据
            form_data = aiohttp.FormData()
            form_data.add_field('file',
                              io.BytesIO(image_data),
                              filename=filename,
                              content_type='image/jpeg')

            timeout = aiohttp.ClientTimeout(total=FileService.IMAGE_UPLOAD_TIMEOUT)
            async with aiohttp.ClientSession(timeout=timeout) as session:
                async with session.post(FileService.IMAGE_UPLOAD_URL, data=form_data) as response:
                    if response.status == 200:
                        result = await response.json()
                        # 根据实际API返回格式获取图片URL
                        return result.get('file_url')
                    else:
                        print(f"图片上传失败，状态码: {response.status}")
                        return None
        except Exception as e:
            print(f"图片上传异常: {str(e)}")
            return None

    @staticmethod
    def extract_images_from_pdf(file_path: str) -> List[Tuple[bytes, str, int, int]]:
        """从PDF提取图片，返回 (图片数据, 扩展名, 页码, 图片索引) 列表"""
        if not HAS_ADVANCED_LIBS:
            return []

        images = []
        try:
            doc = fitz.open(file_path)

            for page_num in range(doc.page_count):
                page = doc[page_num]
                image_list = page.get_images(full=True)

                for img_index, img in enumerate(image_list):
                    try:
                        # 获取图片数据
                        xref = img[0]
                        pix = fitz.Pixmap(doc, xref)

                        # 转换为RGB格式（如果是CMYK）
                        if pix.n - pix.alpha < 4:
                            img_data = pix.tobytes("jpeg")
                            ext = "jpg"
                        else:
                            pix1 = fitz.Pixmap(fitz.csRGB, pix)
                            img_data = pix1.tobytes("jpeg")
                            ext = "jpg"
                            pix1 = None

                        pix = None
                        images.append((img_data, ext, page_num + 1, img_index + 1))

                    except Exception as e:
                        print(f"提取PDF第{page_num+1}页图片{img_index+1}失败: {str(e)}")
                        continue

            doc.close()
            return images

        except Exception as e:
            print(f"PDF图片提取失败: {str(e)}")
            return []

    @staticmethod
    def extract_images_from_docx(file_path: str) -> List[Tuple[bytes, str, int]]:
        """从Word文档提取图片，返回 (图片数据, 扩展名, 图片索引) 列表"""
        images = []
        doc = None
        try:
            doc = docx.Document(file_path)

            # 获取文档中的所有关系
            rels = doc.part.rels
            img_index = 0

            for rel in rels.values():
                if "image" in rel.target_ref:
                    try:
                        # 读取图片数据
                        img_data = rel.target_part.blob

                        # 根据content_type确定扩展名
                        content_type = rel.target_part.content_type
                        if 'jpeg' in content_type:
                            ext = 'jpg'
                        elif 'png' in content_type:
                            ext = 'png'
                        elif 'gif' in content_type:
                            ext = 'gif'
                        elif 'bmp' in content_type:
                            ext = 'bmp'
                        else:
                            ext = 'jpg'  # 默认

                        img_index += 1
                        images.append((img_data, ext, img_index))

                    except Exception as e:
                        print(f"提取Word文档图片{img_index+1}失败: {str(e)}")
                        continue

            if doc:
                del doc
            gc.collect()
            return images

        except Exception as e:
            if doc:
                del doc
            gc.collect()
            print(f"Word文档图片提取失败: {str(e)}")
            return []

    @staticmethod
    def _safe_file_cleanup(file_path: str, max_retries: int = 3) -> bool:
        """安全删除文件，带重试机制"""
        for attempt in range(max_retries):
            try:
                if os.path.exists(file_path):
                    # 强制垃圾回收，释放可能的文件句柄
                    gc.collect()
                    time.sleep(0.1 * (attempt + 1))  # 递增延迟
                    os.remove(file_path)
                return True
            except OSError as e:
                if attempt == max_retries - 1:
                    print(f"无法删除文件 {file_path}: {e}")
                    return False
                time.sleep(0.5)  # 等待后重试
        return True
    
    @staticmethod
    async def save_uploaded_file(file: UploadFile) -> str:
        """保存上传的文件并返回文件路径"""
        # 创建上传目录
        os.makedirs(settings.upload_dir, exist_ok=True)

        # 生成带时间戳的文件名，防止重复
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")[:-3]  # 精确到毫秒
        filename = file.filename or "unknown_file"

        # 分离文件名和扩展名
        name, ext = os.path.splitext(filename)

        # 生成新的文件名：原文件名_时间戳.扩展名
        new_filename = f"{name}_{timestamp}{ext}"
        file_path = os.path.join(settings.upload_dir, new_filename)

        # 异步保存文件
        async with aiofiles.open(file_path, 'wb') as f:
            content = await file.read()
            await f.write(content)

        return file_path
    
    @staticmethod
    async def extract_text_from_pdf(file_path: str) -> str:
        """从PDF文件提取文本，支持表格内容和图片"""
        if HAS_ADVANCED_LIBS:
            return await FileService._extract_pdf_with_pdfplumber(file_path)
        else:
            # 降级到原来的PyPDF2方法
            return FileService._extract_pdf_with_pypdf2(file_path)
    
    @staticmethod
    async def _extract_pdf_with_pdfplumber(file_path: str) -> str:
        """使用pdfplumber提取PDF文本，包含表格和图片（确保及时释放文件句柄）"""
        try:
            extracted_text = []
            image_references = []  # 存储图片引用映射
            global_img_counter = 1

            # 获取PDF文档的所有图片信息，用于后续匹配
            all_images = FileService.extract_images_from_pdf(file_path)
            page_images_map = {}
            for img_data, ext, page_num, img_index in all_images:
                if page_num not in page_images_map:
                    page_images_map[page_num] = []
                page_images_map[page_num].append((img_data, ext, img_index))

            # 使用上下文管理器，避免在Windows上产生文件锁
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    # 添加页码标识
                    extracted_text.append(f"\n--- 第 {page_num} 页 ---\n")

                    # 提取普通文本
                    text = page.extract_text()
                    if text:
                        # 检查文本中是否有图片标记
                        import re
                        img_pattern = r'----.*?(?:image|img|media).*?----'
                        img_matches = list(re.finditer(img_pattern, text, re.IGNORECASE))

                        if img_matches and page_num in page_images_map:
                            # 按顺序处理页面中的图片
                            page_images = page_images_map[page_num]
                            processed_text = text

                            for i, match in enumerate(img_matches):
                                if i < len(page_images):
                                    # 获取对应的图片数据
                                    img_data, ext, img_index = page_images[i]
                                    filename = f"pdf_page{page_num}_img{img_index}.{ext}"

                                    # 上传图片
                                    image_url = await FileService.upload_image_to_server(img_data, filename)

                                    if image_url:
                                        # 替换图片标记
                                        old_mark = match.group()
                                        new_mark = f"[图片{global_img_counter}]"
                                        processed_text = processed_text.replace(old_mark, new_mark, 1)

                                        # 记录图片引用
                                        image_references.append(f"[图片{global_img_counter}]: {image_url}")
                                        global_img_counter += 1

                            extracted_text.append(processed_text)
                        else:
                            extracted_text.append(text)

                    # 提取表格
                    tables = page.extract_tables()
                    for table_num, table in enumerate(tables, 1):
                        extracted_text.append(f"\n[表格 {table_num}]")
                        for row in table:
                            if row:  # 跳过空行
                                # 过滤空值并连接单元格
                                row_text = " | ".join([str(cell) if cell else "" for cell in row])
                                extracted_text.append(row_text)
                        extracted_text.append("[表格结束]\n")

            # 在文档末尾添加图片引用映射
            if image_references:
                extracted_text.append(f"\n\n--- 图片引用 ---")
                extracted_text.extend(image_references)

            result = "\n".join(extracted_text).strip()
            gc.collect()
            return result
        except Exception as e:
            gc.collect()
            # 如果pdfplumber失败，尝试PyMuPDF
            try:
                return await FileService._extract_pdf_with_pymupdf(file_path)
            except Exception:
                raise Exception(f"PDF文件读取失败: {str(e)}")
    
    @staticmethod
    async def _extract_pdf_with_pymupdf(file_path: str) -> str:
        """使用PyMuPDF提取PDF文本和图片"""
        try:
            doc = fitz.open(file_path)
            extracted_text = []
            
            for page_num in range(doc.page_count):
                page = doc[page_num]
                extracted_text.append(f"\n--- 第 {page_num + 1} 页 ---\n")
                
                # 提取文本
                text = page.get_text()
                if text:
                    extracted_text.append(text)
                
                # 尝试提取表格
                try:
                    tables = page.find_tables()
                    for table_num, table in enumerate(tables, 1):
                        extracted_text.append(f"\n[表格 {table_num}]")
                        table_data = table.extract()
                        for row in table_data:
                            if row:
                                row_text = " | ".join([str(cell) if cell else "" for cell in row])
                                extracted_text.append(row_text)
                        extracted_text.append("[表格结束]\n")
                except:
                    # 如果表格提取失败，跳过
                    pass
            
            doc.close()
            return "\n".join(extracted_text).strip()
        except Exception as e:
            raise Exception(f"PDF文件读取失败: {str(e)}")
    
    @staticmethod 
    def _extract_pdf_with_pypdf2(file_path: str) -> str:
        """使用PyPDF2提取PDF文本（原方法）"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text.strip()
        except Exception as e:
            raise Exception(f"PDF文件读取失败: {str(e)}")
    
    @staticmethod
    async def extract_text_from_docx(file_path: str) -> str:
        """从Word文档提取文本，支持表格内容和图片"""
        if HAS_ADVANCED_LIBS:
            return await FileService._extract_docx_with_docx2python(file_path)
        else:
            # 降级到原来的python-docx方法，但增强表格处理
            return await FileService._extract_docx_with_python_docx(file_path)

    @staticmethod
    def _build_mineru_env(device: str) -> Dict[str, str]:
        """构造 MinerU 本地运行环境，不调用云端 API。"""
        env = dict(os.environ)
        env.setdefault("PYTORCH_ENABLE_MPS_FALLBACK", "1")
        if device == "cuda":
            env.setdefault("CUDA_VISIBLE_DEVICES", os.getenv("YIBIAO_CUDA_VISIBLE_DEVICES", "0"))
        elif device == "cpu":
            env["CUDA_VISIBLE_DEVICES"] = ""
        return env

    @staticmethod
    def _find_mineru_output_files(output_dir: str) -> Tuple[Optional[Path], Optional[Path]]:
        """在 MinerU 输出目录中找到主 Markdown 和 content_list JSON。"""
        root = Path(output_dir)
        markdown_files = [path for path in root.rglob("*.md") if path.is_file()]
        content_files = [
            path for path in root.rglob("*.json")
            if path.is_file() and "content_list" in path.name
        ]

        markdown = max(markdown_files, key=lambda path: path.stat().st_size, default=None)
        content_json = max(content_files, key=lambda path: path.stat().st_size, default=None)
        return markdown, content_json

    @staticmethod
    def _load_mineru_block_count(content_json: Optional[Path]) -> int:
        if not content_json:
            return 0
        try:
            data = json.loads(content_json.read_text(encoding="utf-8"))
            if isinstance(data, list):
                return len(data)
            if isinstance(data, dict):
                return len(data.get("content") or data.get("blocks") or [])
        except Exception:
            return 0
        return 0

    @staticmethod
    async def _extract_with_mineru(file_path: str, file_kind: str) -> Tuple[str, Dict[str, Any]]:
        """用本机 MinerU CLI 把文件解析成 Markdown。"""
        mineru_bin = FileService._mineru_command()
        if not mineru_bin:
            raise RuntimeError("未找到 mineru 命令，请先安装 MinerU 或设置 YIBIAO_MINERU_BIN")

        device = FileService._detect_mineru_device()
        backend = os.getenv("YIBIAO_MINERU_BACKEND", "pipeline").strip() or "pipeline"
        lang = os.getenv("YIBIAO_MINERU_LANG", "ch").strip() or "ch"
        timeout = int(os.getenv("YIBIAO_MINERU_TIMEOUT", str(FileService.MINERU_DEFAULT_TIMEOUT)))

        with tempfile.TemporaryDirectory(prefix="yibiao-mineru-") as output_dir:
            command = [
                mineru_bin,
                "-p",
                file_path,
                "-o",
                output_dir,
                "-b",
                backend,
                "-l",
                lang,
            ]
            api_url = os.getenv("YIBIAO_MINERU_API_URL", "").strip()
            if api_url:
                command.extend(["--api-url", api_url])

            proc = await asyncio.create_subprocess_exec(
                *command,
                stdout=asyncio.subprocess.PIPE,
                stderr=asyncio.subprocess.PIPE,
                env=FileService._build_mineru_env(device),
            )
            try:
                stdout_bytes, stderr_bytes = await asyncio.wait_for(proc.communicate(), timeout=timeout)
            except asyncio.TimeoutError as exc:
                proc.kill()
                await proc.communicate()
                raise TimeoutError(f"MinerU 解析超过 {timeout} 秒") from exc

            stdout = stdout_bytes.decode("utf-8", errors="replace")
            stderr = stderr_bytes.decode("utf-8", errors="replace")
            if proc.returncode != 0:
                error_tail = (stderr or stdout)[-1200:]
                raise RuntimeError(f"MinerU 解析失败，退出码 {proc.returncode}: {error_tail}")

            markdown_file, content_json = FileService._find_mineru_output_files(output_dir)
            if not markdown_file:
                raise RuntimeError("MinerU 未生成 Markdown 输出")

            markdown = markdown_file.read_text(encoding="utf-8").strip()
            if not markdown:
                raise RuntimeError("MinerU Markdown 输出为空")

            return markdown, {
                "parser": "mineru",
                "format": "markdown",
                "file_kind": file_kind,
                "device": device,
                "backend": backend,
                "language": lang,
                "api_url": api_url,
                "markdown_file": markdown_file.name,
                "content_list_file": content_json.name if content_json else "",
                "content_block_count": FileService._load_mineru_block_count(content_json),
                "fallback_used": False,
            }

    @staticmethod
    async def _extract_with_legacy_parser(file_path: str, file_kind: str) -> Tuple[str, Dict[str, Any]]:
        """现有 pdfplumber / docx2python 解析路径。"""
        if file_kind == "pdf":
            text = await FileService.extract_text_from_pdf(file_path)
            parser = "pdfplumber" if HAS_ADVANCED_LIBS else "pypdf2"
        elif file_kind == "docx":
            text = await FileService.extract_text_from_docx(file_path)
            parser = "docx2python" if HAS_ADVANCED_LIBS else "python-docx"
        else:
            raise Exception("不支持的文件类型")
        return text, {
            "parser": parser,
            "format": "plain_text",
            "file_kind": file_kind,
            "fallback_used": False,
        }

    @staticmethod
    async def _extract_with_configured_parser(file_path: str, file_kind: str) -> Tuple[str, Dict[str, Any]]:
        """按配置优先使用 MinerU，本机不可用时保留现有解析器。"""
        mode = FileService.get_parser_mode()
        if mode == "legacy":
            return await FileService._extract_with_legacy_parser(file_path, file_kind)

        try:
            return await FileService._extract_with_mineru(file_path, file_kind)
        except Exception as exc:
            if mode == "mineru_strict":
                raise
            text, info = await FileService._extract_with_legacy_parser(file_path, file_kind)
            info.update({
                "fallback_used": True,
                "preferred_parser": "mineru",
                "fallback_reason": str(exc),
            })
            return text, info
    
    @staticmethod
    async def _extract_docx_with_docx2python(file_path: str) -> str:
        """使用docx2python提取Word文档内容和图片（确保及时释放文件句柄）"""
        try:
            extracted_text = []
            image_references = []  # 存储图片引用映射
            global_img_counter = 1

            # 获取Word文档的所有图片信息
            all_images = FileService.extract_images_from_docx(file_path)

            # 使用上下文管理器确保文件及时关闭，避免Windows上的锁定
            with docx2python(file_path) as content:
                # 处理文档内容
                if hasattr(content, 'document'):
                    for section in content.document:
                        for element in section:
                            if isinstance(element, list):
                                # 这可能是表格
                                extracted_text.append("\n[表格内容]")
                                for row in element:
                                    if isinstance(row, list):
                                        row_text = " | ".join([str(cell).strip() for cell in row if cell])
                                        if row_text:
                                            extracted_text.append(row_text)
                                    else:
                                        extracted_text.append(str(row))
                                extracted_text.append("[表格结束]\n")
                            else:
                                # 普通文本，检查是否包含图片标记
                                text = str(element).strip()
                                if text:
                                    # 检查文本中是否有图片标记
                                    import re
                                    img_pattern = r'----.*?(?:image|img|media).*?----'
                                    img_matches = list(re.finditer(img_pattern, text, re.IGNORECASE))

                                    if img_matches and all_images:
                                        processed_text = text

                                        for match in img_matches:
                                            if global_img_counter <= len(all_images):
                                                # 获取对应的图片数据
                                                img_data, ext, img_index = all_images[global_img_counter - 1]
                                                filename = f"docx_img{global_img_counter}.{ext}"

                                                # 上传图片
                                                image_url = await FileService.upload_image_to_server(img_data, filename)

                                                if image_url:
                                                    # 替换图片标记
                                                    old_mark = match.group()
                                                    new_mark = f"[图片{global_img_counter}]"
                                                    processed_text = processed_text.replace(old_mark, new_mark, 1)

                                                    # 记录图片引用
                                                    image_references.append(f"[图片{global_img_counter}]: {image_url}")
                                                    global_img_counter += 1

                                        extracted_text.append(processed_text)
                                    else:
                                        extracted_text.append(text)

            # 在文档末尾添加图片引用映射
            if image_references:
                extracted_text.append(f"\n\n--- 图片引用 ---")
                extracted_text.extend(image_references)

            result = "\n".join(extracted_text).strip()
            gc.collect()
            return result
        except Exception as e:
            gc.collect()
            # 如果docx2python失败，回退到增强的python-docx
            try:
                return await FileService._extract_docx_with_python_docx(file_path)
            except Exception:
                raise Exception(f"Word文档读取失败: {str(e)}")
    
    @staticmethod
    async def _extract_docx_with_python_docx(file_path: str) -> str:
        """使用python-docx提取Word文档内容和图片（增强版）"""
        doc = None
        try:
            doc = docx.Document(file_path)
            extracted_text = []
            image_references = []  # 存储图片引用映射
            global_img_counter = 1

            # 获取Word文档的所有图片信息
            all_images = FileService.extract_images_from_docx(file_path)

            # 提取段落文本，同时处理图片
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    # 检查文本中是否有图片标记
                    import re
                    img_pattern = r'----.*?(?:image|img|media).*?----'
                    img_matches = list(re.finditer(img_pattern, text, re.IGNORECASE))

                    if img_matches and all_images:
                        processed_text = text

                        for match in img_matches:
                            if global_img_counter <= len(all_images):
                                # 获取对应的图片数据
                                img_data, ext, img_index = all_images[global_img_counter - 1]
                                filename = f"docx_img{global_img_counter}.{ext}"

                                # 上传图片
                                image_url = await FileService.upload_image_to_server(img_data, filename)

                                if image_url:
                                    # 替换图片标记
                                    old_mark = match.group()
                                    new_mark = f"[图片{global_img_counter}]"
                                    processed_text = processed_text.replace(old_mark, new_mark, 1)

                                    # 记录图片引用
                                    image_references.append(f"[图片{global_img_counter}]: {image_url}")
                                    global_img_counter += 1

                        extracted_text.append(processed_text)
                    else:
                        extracted_text.append(text)

            # 提取表格内容
            for table_num, table in enumerate(doc.tables, 1):
                extracted_text.append(f"\n[表格 {table_num}]")
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        row_data.append(cell_text if cell_text else "")
                    row_text = " | ".join(row_data)
                    if row_text.strip():
                        extracted_text.append(row_text)
                extracted_text.append("[表格结束]\n")

            # 在文档末尾添加图片引用映射
            if image_references:
                extracted_text.append(f"\n\n--- 图片引用 ---")
                extracted_text.extend(image_references)

            result = "\n".join(extracted_text).strip()

            # 确保释放资源
            if doc:
                del doc
            gc.collect()

            return result
        except Exception as e:
            # 确保释放资源
            if doc:
                del doc
            gc.collect()
            raise Exception(f"Word文档读取失败: {str(e)}")
    
    @staticmethod
    async def process_uploaded_file_with_metadata(file: UploadFile) -> Dict[str, Any]:
        """处理上传的文件并提取文本内容，返回解析器元信息。"""
        file_kind = FileService.detect_upload_file_kind(file)
        if file_kind in (None, "doc"):
            raise Exception(FileService.get_upload_validation_message(file))

        # 检查文件大小
        content = await file.read()
        if len(content) > settings.max_file_size:
            raise Exception(f"文件大小超过限制 ({settings.max_file_size / 1024 / 1024}MB)")
        
        # 重置文件指针
        await file.seek(0)
        
        # 保存文件
        file_path = await FileService.save_uploaded_file(file)
        
        try:
            text, parser_info = await FileService._extract_with_configured_parser(file_path, file_kind)
            if not text.strip():
                raise Exception("无法从文件中提取文本内容")
            parser_info.setdefault("saved_file", os.path.basename(file_path))
            parser_info.setdefault("file_size", os.path.getsize(file_path))

            # 成功提取后，使用安全的文件清理方法
            FileService._safe_file_cleanup(file_path)

            return {
                "file_content": text,
                "parser_info": parser_info,
            }

        except Exception as e:
            # 异常情况下也使用安全的文件清理方法
            FileService._safe_file_cleanup(file_path)
            raise e

    @staticmethod
    async def process_uploaded_file(file: UploadFile) -> str:
        """处理上传的文件并提取文本内容。"""
        result = await FileService.process_uploaded_file_with_metadata(file)
        return str(result.get("file_content") or "")
