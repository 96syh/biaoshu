"""Word export service for bid documents."""
import base64
import io
import json
import re
import urllib.request
from pathlib import Path
from urllib.parse import quote, unquote

import docx
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from fastapi import HTTPException
from fastapi.responses import JSONResponse, StreamingResponse

from ..models.schemas import WordExportRequest
from .file_service import FileService
from .generation.content import ContentGenerationMixin
from .history_case_service import HistoryCaseService


_INHERITED_TEMPLATE_MODE = False


def sanitize_docx_filename(filename: str) -> str:
    """清理保存到本地文件系统的 docx 文件名。"""
    safe = re.sub(r'[\\/:*?"<>|\r\n]+', "_", filename or "标书文档.docx").strip()
    if not safe.lower().endswith(".docx"):
        safe = f"{safe}.docx"
    return safe or "标书文档.docx"


def set_run_font_simsun(run: docx.text.run.Run) -> None:
    """统一将 run 字体设置为宋体（包含 EastAsia 字体设置）"""
    if _INHERITED_TEMPLATE_MODE:
        return
    run.font.name = "宋体"
    r = run._element.rPr
    if r is not None and r.rFonts is not None:
        r.rFonts.set(qn("w:eastAsia"), "宋体")


def set_paragraph_font_simsun(paragraph: docx.text.paragraph.Paragraph) -> None:
    """将段落内所有 runs 字体设置为宋体"""
    if _INHERITED_TEMPLATE_MODE:
        return
    for run in paragraph.runs:
        set_run_font_simsun(run)


def add_field(paragraph: docx.text.paragraph.Paragraph, instruction: str) -> None:
    """Add a Word field that can be updated inside Word/WPS."""
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "请在 Word 中更新域"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    run = paragraph.add_run()
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(placeholder)
    run._r.append(end)
    set_run_font_simsun(run)


def enable_update_fields_on_open(doc: docx.Document) -> None:
    """Ask Word to update TOC/page-number fields when opening the document."""
    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
    settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def _style_value(style_profile: dict, key: str, default: str) -> str:
    value = style_profile.get(key)
    return str(value).strip() if value else default


def _cm_value(value: str, default: float) -> float:
    text = str(value or "").strip().lower().replace("厘米", "cm")
    try:
        if text.endswith("cm"):
            return float(text[:-2])
        if text.endswith("mm"):
            return float(text[:-2]) / 10
        return float(text)
    except Exception:
        return default


def _pt_value(value: str, default: float) -> float:
    text = str(value or "").strip().lower().replace("小四", "12pt").replace("五号", "10.5pt")
    try:
        if text.endswith("pt"):
            return float(text[:-2])
        return float(text)
    except Exception:
        return default


def configure_word_style(doc: docx.Document, reference_profile: dict | None = None) -> None:
    """Apply mature-sample-like Word page and font settings."""
    profile = reference_profile or {}
    style_profile = profile.get("word_style_profile") if isinstance(profile, dict) else {}
    if not isinstance(style_profile, dict):
        style_profile = {}

    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(_cm_value(_style_value(style_profile, "margin_top", "2.2cm"), 2.2))
        section.bottom_margin = Cm(_cm_value(_style_value(style_profile, "margin_bottom", "2.2cm"), 2.2))
        section.left_margin = Cm(_cm_value(_style_value(style_profile, "margin_left", "2.7cm"), 2.7))
        section.right_margin = Cm(_cm_value(_style_value(style_profile, "margin_right", "2.2cm"), 2.2))
        section.header_distance = Cm(1.5)
        section.footer_distance = Cm(1.75)

    body_font = _style_value(style_profile, "body_font_family", "宋体").split(",")[0].strip().strip('"')
    heading_font = _style_value(style_profile, "heading_font_family", "黑体").split(",")[0].strip().strip('"')
    body_size = _pt_value(_style_value(style_profile, "body_font_size", "10.5pt"), 10.5)

    def set_style_font(style_name: str, font_name: str, size_pt: float, bold: bool | None = None) -> None:
        style = doc.styles[style_name]
        style.font.name = font_name
        style.font.size = Pt(size_pt)
        if bold is not None:
            style.font.bold = bold
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

    set_style_font("Normal", body_font, body_size, False)
    set_style_font("Heading 1", heading_font, _pt_value(_style_value(style_profile, "heading_1_size", "16pt"), 16), True)
    set_style_font("Heading 2", heading_font, _pt_value(_style_value(style_profile, "heading_2_size", "14pt"), 14), True)
    set_style_font("Heading 3", heading_font, _pt_value(_style_value(style_profile, "heading_3_size", "12pt"), 12), True)


def clear_document_body_keep_section(doc: docx.Document) -> None:
    """Clear template body while preserving styles, section settings, headers and footers."""
    body = doc._element.body
    sect_pr = body.sectPr
    for child in list(body):
        if sect_pr is not None and child is sect_pr:
            continue
        body.remove(child)
    if sect_pr is not None and sect_pr.getparent() is None:
        body.append(sect_pr)


def iter_outline_items(items) -> list:
    result = []
    for item in items or []:
        result.append(item)
        result.extend(iter_outline_items(getattr(item, "children", None) or []))
    return result


def resolve_history_template_docx(request: WordExportRequest) -> Path | None:
    """Pick the first matched historical docx as the export style template."""
    for item in iter_outline_items(request.outline):
        reference = getattr(item, "history_reference", None) or {}
        if not isinstance(reference, dict):
            continue
        source_paths = reference.get("source_paths") or reference
        if not isinstance(source_paths, dict):
            continue
        raw_path = str(
            source_paths.get("source_docx_path")
            or source_paths.get("best_document_path")
            or ""
        ).strip()
        if not raw_path:
            continue
        path = Path(raw_path).expanduser()
        if path.exists() and path.is_file() and path.suffix.lower() == ".docx":
            return path
    return None


def create_export_document(request: WordExportRequest) -> tuple[docx.Document, Path | None]:
    template_path = resolve_history_template_docx(request)
    if template_path:
        doc = docx.Document(str(template_path))
        clear_document_body_keep_section(doc)
        return doc, template_path
    return docx.Document(), None



async def create_word_export_response(request: WordExportRequest):
    """根据目录数据导出Word文档"""
    global _INHERITED_TEMPLATE_MODE
    previous_template_mode = _INHERITED_TEMPLATE_MODE
    try:
        if not request.manual_review_confirmed:
            raise HTTPException(status_code=400, detail="导出前必须完成人工复核确认，不能直接跳过模型结果复核。")

        doc, inherited_template_path = create_export_document(request)
        _INHERITED_TEMPLATE_MODE = bool(inherited_template_path)
        enable_update_fields_on_open(doc)
        if not inherited_template_path:
            configure_word_style(doc, request.reference_bid_style_profile)

        # 统一设置普通正文基础字体，标题样式已由成熟样例模板配置。
        if not inherited_template_path:
            try:
                styles = doc.styles
                base_styles = ["Normal"]
                for style_name in base_styles:
                    if style_name in styles:
                        style = styles[style_name]
                        font = style.font
                        font.name = "宋体"
                        # 设置中文字体
                        if style._element.rPr is None:
                            style._element._add_rPr()
                        rpr = style._element.rPr
                        rpr.rFonts.set(qn("w:eastAsia"), "宋体")
                        if style_name == "Normal":
                            font.bold = False
            except Exception:
                # 字体设置失败不影响文档生成，忽略
                pass

        # AI 生成声明
        p = doc.add_paragraph()
        run = p.add_run("内容由 AI 生成，导出前已要求人工复核确认；最终页码、目录、签章、版式和图表需在 Word 中复核。")
        run.italic = True
        if not inherited_template_path:
            run.font.size = Pt(9)
        set_run_font_simsun(run)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        footer_p = doc.sections[0].footer.paragraphs[0]
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_p.add_run("第 ")
        add_field(footer_p, "PAGE")
        footer_p.add_run(" 页 / 共 ")
        add_field(footer_p, "NUMPAGES")
        footer_p.add_run(" 页")
        set_paragraph_font_simsun(footer_p)

        # 文档标题
        title = request.project_name or "投标技术文件"
        title_p = doc.add_paragraph()
        title_run = title_p.add_run(title)
        if not inherited_template_path:
            title_run.bold = True
        if not inherited_template_path:
            title_run.font.size = Pt(16)
        set_run_font_simsun(title_run)
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 可更新目录。python-docx 不能计算最终页码，但可以生成 Word 域。
        toc_heading = doc.add_heading("目录", level=1)
        set_paragraph_font_simsun(toc_heading)
        toc_p = doc.add_paragraph()
        add_field(toc_p, r'TOC \o "1-3" \h \z \u')
        doc.add_page_break()

        def add_export_checklist() -> None:
            """Add a visible finalization checklist for manual Word work."""
            report = request.analysis_report.model_dump(mode="json") if request.analysis_report else {}
            review = request.review_report.model_dump(mode="json") if request.review_report else {}
            bid_rules = report.get("bid_document_requirements") or {}
            formatting = bid_rules.get("formatting_and_submission_rules") or {}
            signature_items = report.get("signature_requirements") or []
            fixed_forms = report.get("fixed_format_forms") or []
            enterprise_profile = report.get("enterprise_material_profile") or {}
            missing_enterprise = enterprise_profile.get("missing_materials") or report.get("missing_company_materials") or []
            plan = request.document_blocks_plan or {}
            missing_assets = []
            if isinstance(plan, dict):
                missing_assets = (plan.get("missing_assets") or []) + (plan.get("missing_enterprise_data") or [])
            blocking_count = ((review.get("summary") or {}).get("blocking_issues_count")
                              or (review.get("summary") or {}).get("blocking_issues") or 0)

            heading = doc.add_heading("导出后 Word 处理清单", level=1)
            set_paragraph_font_simsun(heading)
            lines = [
                "更新目录域和页码域：打开 Word 后全选并更新域，确认目录层级和页码正确。",
                "复核版式：按招标文件检查页边距、字体、行距、标题层级、表格跨页和页眉页脚。",
                "复核签章位置：按签章要求在投标函、授权委托书、报价表、承诺函、固定格式表单等位置签字盖章。",
                "复核固定格式：不得破坏招标文件固定表头、固定文字、行列结构和附件说明。",
                "复核图表与附件：将图表、组织架构图、流程图、证书、截图、扫描件替换到对应占位。",
                "复核模型风险：逐项核对误读、漏项、虚构、历史项目残留和格式偏差；不能仅凭模型审校结论提交。",
            ]
            if formatting.get("toc_required") or formatting.get("page_number_required"):
                lines.append("招标文件要求目录/页码，请在最终版中更新目录和所有响应页码。")
            if signature_items:
                lines.append(f"已解析签章要求 {len(signature_items)} 项，需逐项核验签署主体、盖章、日期和电子签章。")
            if fixed_forms:
                lines.append(f"已解析固定格式 {len(fixed_forms)} 项，需人工确认表头、固定文字和签章栏。")
            if missing_enterprise:
                lines.append(f"企业资料仍有 {len(missing_enterprise)} 项待补或待确认，正文中的待补占位不得直接提交。")
            if missing_assets:
                lines.append(f"图表/素材/企业数据仍有 {len(missing_assets)} 项待替换。")
            if blocking_count:
                lines.append(f"审校报告仍有 {blocking_count} 项阻塞问题，处理后再提交最终版。")

            for line in lines:
                item_p = doc.add_paragraph()
                run = item_p.add_run(f"□ {line}")
                set_run_font_simsun(run)
            doc.add_page_break()

        add_export_checklist()

        # 项目概述
        if request.project_overview:
            heading = doc.add_heading("项目概述", level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_font_simsun(heading)
            overview_p = doc.add_paragraph(request.project_overview)
            set_paragraph_font_simsun(overview_p)
            overview_p_format = overview_p.paragraph_format
            overview_p_format.space_after = Pt(12)

        # 简单的 Markdown 段落解析：支持标题、列表、表格和基础加粗/斜体
        def add_markdown_runs(para: docx.text.paragraph.Paragraph, text: str) -> None:
            """在指定段落中追加 markdown 文本的 runs"""
            pattern = r"(\*\*.*?\*\*|\*.*?\*|`.*?`)"
            parts = re.split(pattern, text)
            for part in parts:
                if not part:
                    continue
                run = para.add_run()
                # 加粗
                if part.startswith("**") and part.endswith("**") and len(part) > 4:
                    run.text = part[2:-2]
                    run.bold = True
                # 斜体
                elif part.startswith("*") and part.endswith("*") and len(part) > 2:
                    run.text = part[1:-1]
                    run.italic = True
                # 行内代码：这里只去掉反引号
                elif part.startswith("`") and part.endswith("`") and len(part) > 2:
                    run.text = part[1:-1]
                else:
                    run.text = part
                # 确保字体为宋体
                set_run_font_simsun(run)

        def add_markdown_paragraph(text: str) -> None:
            """将一段 Markdown 文本解析为一个普通段落，保留加粗/斜体效果"""
            para = doc.add_paragraph()
            add_markdown_runs(para, text)
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.first_line_indent = Pt(21)
            para.paragraph_format.line_spacing = 1.5

        def parse_markdown_table_row(text: str) -> list[str]:
            """Parse a Markdown pipe table row into cells."""
            stripped = text.strip()
            if "|" not in stripped:
                return []
            if stripped.startswith("|"):
                stripped = stripped[1:]
            if stripped.endswith("|"):
                stripped = stripped[:-1]
            cells = [re.sub(r"<br\s*/?>", "\n", cell.strip(), flags=re.IGNORECASE) for cell in stripped.split("|")]
            if len(cells) < 2:
                return []
            return cells

        def is_markdown_table_separator(text: str) -> bool:
            cells = parse_markdown_table_row(text)
            if not cells:
                return False
            return all(re.match(r"^:?-{3,}:?$", cell.replace(" ", "")) for cell in cells)

        def normalize_table_rows(rows: list[list[str]]) -> list[list[str]]:
            max_cols = max((len(row) for row in rows), default=0)
            return [row + [""] * (max_cols - len(row)) for row in rows if any(cell.strip() for cell in row)]

        def content_width_dxa() -> int:
            section = doc.sections[-1]
            width_emu = int(section.page_width) - int(section.left_margin) - int(section.right_margin)
            return max(3600, int(width_emu / 635))

        def table_column_widths(rows: list[list[str]], table_width_dxa: int) -> list[int]:
            cols = max((len(row) for row in rows), default=0)
            if cols <= 0:
                return []
            weights = []
            for col_idx in range(cols):
                col_texts = [row[col_idx] if col_idx < len(row) else "" for row in rows]
                max_len = max((len(re.sub(r"\s+", "", str(text))) for text in col_texts), default=1)
                weights.append(max(1.0, min(4.0, max_len / 7)))
            total_weight = sum(weights) or cols
            widths = [max(240, int(round(table_width_dxa * weight / total_weight))) for weight in weights]
            widths[-1] += table_width_dxa - sum(widths)
            return widths

        def get_or_add(parent, tag: str):
            child = parent.find(qn(tag))
            if child is None:
                child = OxmlElement(tag)
                parent.append(child)
            return child

        def set_table_property(table, table_width_dxa: int, widths: list[int]) -> None:
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            table.autofit = False
            tbl = table._tbl
            tbl_pr = tbl.tblPr

            tbl_w = get_or_add(tbl_pr, "w:tblW")
            tbl_w.set(qn("w:type"), "dxa")
            tbl_w.set(qn("w:w"), str(table_width_dxa))

            tbl_ind = get_or_add(tbl_pr, "w:tblInd")
            tbl_ind.set(qn("w:type"), "dxa")
            tbl_ind.set(qn("w:w"), "0")

            tbl_layout = get_or_add(tbl_pr, "w:tblLayout")
            tbl_layout.set(qn("w:type"), "fixed")

            tbl_borders = get_or_add(tbl_pr, "w:tblBorders")
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                element = get_or_add(tbl_borders, f"w:{edge}")
                element.set(qn("w:val"), "nil")
                element.set(qn("w:sz"), "0")
                element.set(qn("w:space"), "0")

            existing_grid = tbl.tblGrid
            if existing_grid is not None:
                tbl.remove(existing_grid)
            tbl_grid = OxmlElement("w:tblGrid")
            for width in widths:
                grid_col = OxmlElement("w:gridCol")
                grid_col.set(qn("w:w"), str(width))
                tbl_grid.append(grid_col)
            tbl.insert(1, tbl_grid)

        def set_cell_border(cell, **edges: dict[str, str]) -> None:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_borders = get_or_add(tc_pr, "w:tcBorders")
            for edge, attrs in edges.items():
                element = get_or_add(tc_borders, f"w:{edge}")
                for key, value in attrs.items():
                    element.set(qn(f"w:{key}"), str(value))

        def set_cell_margins_and_width(cell, width_dxa: int) -> None:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = get_or_add(tc_pr, "w:tcW")
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width_dxa))

            tc_mar = get_or_add(tc_pr, "w:tcMar")
            for side, value in (("top", 80), ("bottom", 80), ("left", 120), ("right", 120)):
                margin = get_or_add(tc_mar, f"w:{side}")
                margin.set(qn("w:w"), str(value))
                margin.set(qn("w:type"), "dxa")

        def set_table_cell_text(cell, text: str, is_header: bool, column_idx: int) -> None:
            cell.text = ""
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            lines = str(text or "").splitlines() or [""]
            for line_idx, line in enumerate(lines):
                para = cell.paragraphs[0] if line_idx == 0 else cell.add_paragraph()
                add_markdown_runs(para, line.strip())
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)
                para.paragraph_format.line_spacing = 1.25
                compact_text = re.sub(r"\s+", "", line)
                if is_header or column_idx == 0 or len(compact_text) <= 12:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in para.runs:
                    if is_header and not _INHERITED_TEMPLATE_MODE:
                        run.bold = True
                    if not _INHERITED_TEMPLATE_MODE:
                        run.font.size = Pt(10.5)
                    set_run_font_simsun(run)

        def add_three_line_table(rows: list[list[str]]) -> None:
            normalized_rows = normalize_table_rows(rows)
            if not normalized_rows:
                return
            cols = len(normalized_rows[0])
            table_width = content_width_dxa()
            widths = table_column_widths(normalized_rows, table_width)
            table = doc.add_table(rows=len(normalized_rows), cols=cols)
            set_table_property(table, table_width, widths)

            empty_border = {"val": "nil", "sz": "0", "space": "0"}
            strong_line = {"val": "single", "sz": "12", "space": "0", "color": "000000"}
            header_line = {"val": "single", "sz": "8", "space": "0", "color": "000000"}
            last_row_idx = len(normalized_rows) - 1

            for row_idx, row in enumerate(normalized_rows):
                if row_idx == 0:
                    tr_pr = table.rows[row_idx]._tr.get_or_add_trPr()
                    if tr_pr.find(qn("w:tblHeader")) is None:
                        tbl_header = OxmlElement("w:tblHeader")
                        tbl_header.set(qn("w:val"), "true")
                        tr_pr.append(tbl_header)

                for col_idx, value in enumerate(row):
                    cell = table.rows[row_idx].cells[col_idx]
                    set_cell_margins_and_width(cell, widths[col_idx])
                    set_table_cell_text(cell, value, row_idx == 0, col_idx)

                    borders = {
                        "top": empty_border,
                        "left": empty_border,
                        "bottom": empty_border,
                        "right": empty_border,
                    }
                    if row_idx == 0:
                        borders["top"] = strong_line
                        borders["bottom"] = header_line if last_row_idx > 0 else strong_line
                    if row_idx == last_row_idx:
                        borders["bottom"] = strong_line
                    set_cell_border(cell, **borders)

            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(4)

        def parse_markdown_blocks(content: str):
            """
            识别 Markdown 内容中的块级元素，返回结构化的 block 列表：
            - ('list', items)        items: [(kind, num_str, text), ...]
            - ('table', rows)        rows: [[cell, ...], ...]
            - ('heading', level, text)
            - ('image', alt, src)
            - ('paragraph', text)
            """
            blocks = []
            lines = content.split("\n")
            i = 0
            while i < len(lines):
                line = lines[i].rstrip("\r").strip()
                if not line:
                    i += 1
                    continue

                image_match = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$", line)
                if image_match:
                    blocks.append(("image", image_match.group(1).strip(), image_match.group(2).strip()))
                    i += 1
                    continue

                # 列表项（有序/无序）
                if line.startswith("- ") or line.startswith("* ") or re.match(r"^\d+\.\s", line):
                    # items: (kind, number, text)
                    items = []
                    while i < len(lines):
                        raw = lines[i].rstrip("\r")
                        stripped = raw.strip()
                        # 无序列表
                        if stripped.startswith("- ") or stripped.startswith("* "):
                            text = re.sub(r"^[-*]\s+", "", stripped).strip()
                            if text:
                                items.append(("unordered", None, text))
                            i += 1
                            continue
                        # 有序列表（1. xxx）
                        m_num = re.match(r"^(\d+)\.\s+(.*)$", stripped)
                        if m_num:
                            num_str, text = m_num.groups()
                            text = text.strip()
                            if text:
                                items.append(("ordered", num_str, text))
                            i += 1
                            continue
                        break

                    if items:
                        blocks.append(("list", items))
                    continue

                # 表格：输出为 Word 真实表格，导出时渲染为三线表
                if parse_markdown_table_row(line):
                    rows = []
                    while i < len(lines):
                        raw = lines[i].rstrip("\r")
                        stripped = raw.strip()
                        cells = parse_markdown_table_row(stripped)
                        if not cells:
                            break
                        if not is_markdown_table_separator(stripped):
                            rows.append(cells)
                        i += 1
                    if rows:
                        blocks.append(("table", normalize_table_rows(rows)))
                    continue

                # Markdown 标题（# / ## / ###）
                if line.startswith("#"):
                    m = re.match(r"^(#+)\s*(.*)$", line)
                    if m:
                        level_marks, title_text = m.groups()
                        level = min(len(level_marks), 3)
                        blocks.append(("heading", level, title_text.strip()))
                    i += 1
                    continue

                # 普通段落：合并连续的普通行
                para_lines = []
                while i < len(lines):
                    raw = lines[i].rstrip("\r")
                    stripped = raw.strip()
                    if (
                        stripped
                        and not stripped.startswith("-")
                        and not stripped.startswith("*")
                        and not parse_markdown_table_row(stripped)
                        and not stripped.startswith("#")
                    ):
                        para_lines.append(stripped)
                        i += 1
                    else:
                        break
                if para_lines:
                    text = " ".join(para_lines)
                    blocks.append(("paragraph", text))
                else:
                    i += 1

            return blocks

        def resolve_history_markdown_image(src: str, item) -> bytes | None:
            source_paths = getattr(item, "history_reference", {}) or {}
            if isinstance(source_paths, dict):
                source_paths = source_paths.get("source_paths") or source_paths
            asset_dir = str(source_paths.get("asset_dir") or "") if isinstance(source_paths, dict) else ""
            if not asset_dir:
                return None
            relative = str(src or "").replace("\\", "/")
            if relative.startswith("assets/"):
                relative = relative.split("/", 1)[1]
            candidate = (Path(asset_dir) / relative).resolve()
            try:
                asset_root = Path(asset_dir).resolve()
                if asset_root in candidate.parents and candidate.exists() and candidate.is_file():
                    return candidate.read_bytes()
            except Exception:
                return None
            return None

        def render_markdown_blocks(blocks, item=None) -> None:
            """将结构化的 Markdown blocks 渲染到文档"""
            for block in blocks:
                kind = block[0]
                if kind == "list":
                    items = block[1]
                    for item_kind, num_str, text in items:
                        p = doc.add_paragraph()
                        if item_kind == "unordered":
                            # 使用“• ”模拟项目符号
                            run = p.add_run("• ")
                            set_run_font_simsun(run)
                        else:
                            # 有序列表：输出 "1. " 这样的前缀
                            prefix = f"{num_str}."
                            run = p.add_run(prefix + " ")
                            set_run_font_simsun(run)
                        # 紧跟在同一段落中追加列表文本
                        add_markdown_runs(p, text)
                        p.paragraph_format.space_after = Pt(6)
                        p.paragraph_format.line_spacing = 1.5
                elif kind == "table":
                    rows = block[1]
                    add_three_line_table(rows)
                elif kind == "heading":
                    _, level, text = block
                    heading = doc.add_heading(text, level=level)
                    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    set_paragraph_font_simsun(heading)
                elif kind == "image":
                    _, alt, src = block
                    image_bytes = resolve_history_markdown_image(src, item)
                    if image_bytes:
                        try:
                            picture = doc.add_picture(io.BytesIO(image_bytes))
                            max_width = doc.sections[-1].page_width - doc.sections[-1].left_margin - doc.sections[-1].right_margin
                            if picture.width > max_width:
                                ratio = max_width / picture.width
                                picture.width = int(picture.width * ratio)
                                picture.height = int(picture.height * ratio)
                            if alt:
                                caption = doc.add_paragraph(alt)
                                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                set_paragraph_font_simsun(caption)
                        except Exception:
                            add_markdown_paragraph(f"〖图片待插入：{alt or src}〗")
                    else:
                        add_markdown_paragraph(f"〖图片待插入：{alt or src}〗")
                elif kind == "paragraph":
                    _, text = block
                    add_markdown_paragraph(text)

        def add_markdown_content(content: str, item=None) -> None:
            """解析并渲染 Markdown 文本到文档"""
            blocks = parse_markdown_blocks(content)
            render_markdown_blocks(blocks, item=item)

        def item_history_source_paths(item) -> dict:
            reference = getattr(item, "history_reference", {}) or {}
            if not isinstance(reference, dict):
                return {}
            source_paths = reference.get("source_paths") or reference
            return source_paths if isinstance(source_paths, dict) else {}

        def same_docx_path(path_text: str, template_path: Path | None) -> bool:
            if not path_text or not template_path:
                return False
            try:
                return Path(path_text).expanduser().resolve() == template_path.expanduser().resolve()
            except Exception:
                return False

        def append_ooxml_block(xml_text: str) -> bool:
            try:
                element = parse_xml(xml_text)
                body = doc._element.body
                sect_pr = body.sectPr
                if sect_pr is not None:
                    body.insert(list(body).index(sect_pr), element)
                else:
                    body.append(element)
                return True
            except Exception:
                return False

        def remove_duplicate_history_section_heading(blocks: list[dict], item) -> list[dict]:
            if not blocks:
                return blocks
            first = blocks[0]
            if str(first.get("type") or "") != "heading":
                return blocks
            item_title = re.sub(r"\s+", "", str(getattr(item, "title", "") or ""))
            first_title = re.sub(r"\s+", "", str(first.get("text") or ""))
            if item_title and (item_title in first_title or first_title in item_title):
                return blocks[1:]
            return blocks

        def load_patched_history_ooxml_blocks(item) -> list[dict]:
            source_paths = item_history_source_paths(item)
            if not same_docx_path(str(source_paths.get("source_docx_path") or ""), inherited_template_path):
                return []
            reference = getattr(item, "history_reference", {}) or {}
            matched_block_ids = []
            if isinstance(reference, dict) and isinstance(reference.get("matched_block_ids"), list):
                matched_block_ids = [str(block_id) for block_id in reference.get("matched_block_ids") if str(block_id)]
            block_json_path = str(source_paths.get("block_json_path") or "").strip()
            if not block_json_path:
                return []
            try:
                payload = json.loads(Path(block_json_path).expanduser().read_text(encoding="utf-8"))
            except Exception:
                return []
            all_blocks = payload.get("blocks") if isinstance(payload, dict) else []
            if not isinstance(all_blocks, list):
                return []
            if matched_block_ids:
                wanted = set(matched_block_ids)
                matched_blocks = [block for block in all_blocks if str(block.get("id") or "") in wanted]
            else:
                title = str(getattr(item, "title", "") or "")
                matched_blocks = HistoryCaseService._extract_matching_blocks(all_blocks, title)
            matched_blocks = remove_duplicate_history_section_heading(matched_blocks, item)
            if not matched_blocks:
                return []
            operations = getattr(item, "patch_operations", None) or []
            if not isinstance(operations, list):
                operations = []
            return ContentGenerationMixin._apply_history_patch_to_blocks(matched_blocks, operations)

        def add_history_word_blocks_if_possible(item) -> bool:
            if not inherited_template_path:
                return False
            blocks = load_patched_history_ooxml_blocks(item)
            if not blocks:
                return False
            inserted = False
            for block in blocks:
                xml_text = str(block.get("docx_xml") or "").strip()
                if xml_text and append_ooxml_block(xml_text):
                    inserted = True
                else:
                    fallback_text = str(block.get("markdown") or block.get("text") or "").strip()
                    if fallback_text:
                        add_markdown_content(fallback_text, item=item)
                        inserted = True
                caption = str(block.get("caption_text") or "").strip()
                if caption:
                    add_markdown_paragraph(caption)
            return inserted

        def planned_blocks_by_chapter() -> dict[str, list[dict]]:
            """按章节聚合图表/素材规划，兼容分组结构和旧版扁平结构。"""
            plan = request.document_blocks_plan or {}
            groups: dict[str, list[dict]] = {}
            raw_blocks = plan.get("document_blocks") if isinstance(plan, dict) else []
            for group_or_block in raw_blocks or []:
                if not isinstance(group_or_block, dict):
                    continue
                nested_blocks = group_or_block.get("blocks")
                if isinstance(nested_blocks, list):
                    chapter_id = str(group_or_block.get("chapter_id") or group_or_block.get("target_chapter_id") or "").strip()
                    if not chapter_id:
                        continue
                    for block in nested_blocks:
                        if isinstance(block, dict):
                            next_block = dict(block)
                            next_block.setdefault("chapter_id", chapter_id)
                            next_block.setdefault("chapter_title", group_or_block.get("chapter_title") or "")
                            groups.setdefault(chapter_id, []).append(next_block)
                    continue

                chapter_id = str(group_or_block.get("chapter_id") or group_or_block.get("target_chapter_id") or "").strip()
                if chapter_id:
                    groups.setdefault(chapter_id, []).append(group_or_block)
            return groups

        blocks_by_chapter = planned_blocks_by_chapter()
        visual_block_types = {"org_chart", "workflow_chart", "image", "material_attachment"}

        def asset_library_index() -> dict[str, dict]:
            library = request.asset_library or {}
            assets = library.get("visual_assets") if isinstance(library, dict) else []
            index: dict[str, dict] = {}
            if not isinstance(assets, list):
                return index
            for asset in assets:
                if not isinstance(asset, dict):
                    continue
                for key in (
                    asset.get("asset_key"),
                    asset.get("block_key"),
                    f"{asset.get('chapter_id') or ''}:{asset.get('block_name') or ''}",
                ):
                    key_text = str(key or "").strip()
                    if key_text:
                        index[key_text] = asset
            return index

        asset_index = asset_library_index()

        def generated_asset_for_block(block: dict, chapter_id: str, block_name: str) -> dict:
            for key in ("generated_asset", "visual_asset", "asset"):
                value = block.get(key)
                if isinstance(value, dict) and (value.get("image_url") or value.get("b64_json") or value.get("url")):
                    return value
            for key in (
                block.get("asset_key"),
                block.get("block_key"),
                f"{chapter_id}:{block_name}",
            ):
                asset = asset_index.get(str(key or "").strip())
                if asset:
                    return asset
            return {}

        def decode_b64_payload(value: str) -> bytes | None:
            payload = (value or "").strip()
            if not payload:
                return None
            if payload.startswith("data:image") and "," in payload:
                payload = payload.split(",", 1)[1]
            try:
                return base64.b64decode(payload)
            except Exception:
                return None

        def resolve_image_bytes(asset: dict) -> bytes | None:
            b64_json = str(asset.get("b64_json") or asset.get("base64") or asset.get("image_base64") or "").strip()
            decoded = decode_b64_payload(b64_json)
            if decoded:
                return decoded

            image_url = str(asset.get("image_url") or asset.get("url") or asset.get("file_url") or "").strip()
            decoded = decode_b64_payload(image_url)
            if decoded:
                return decoded
            if image_url.startswith(FileService.GENERATED_ASSET_URL_PREFIX + "/"):
                relative_name = unquote(image_url.rsplit("/", 1)[-1])
                candidate = (FileService.GENERATED_ASSET_DIR / relative_name).resolve()
                root = FileService.GENERATED_ASSET_DIR.resolve()
                if str(candidate).startswith(str(root)) and candidate.exists() and candidate.is_file():
                    return candidate.read_bytes()
            if image_url.startswith("http://") or image_url.startswith("https://"):
                try:
                    req = urllib.request.Request(image_url, headers={"User-Agent": "yibiao-export/1.0"})
                    with urllib.request.urlopen(req, timeout=20) as response:
                        return response.read(12 * 1024 * 1024)
                except Exception:
                    return None
            return None

        def add_generated_visual_asset(asset: dict, block_name: str, chapter_id: str, figure_index: int) -> bool:
            image_bytes = resolve_image_bytes(asset)
            if not image_bytes:
                return False
            try:
                picture = doc.add_picture(io.BytesIO(image_bytes))
                section = doc.sections[-1]
                max_width = section.page_width - section.left_margin - section.right_margin
                if picture.width > max_width:
                    ratio = max_width / picture.width
                    picture.width = max_width
                    picture.height = int(picture.height * ratio)
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_text = str(asset.get("caption") or f"图 {chapter_id}-{figure_index} {block_name}").strip()
                caption = doc.add_paragraph()
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = caption.add_run(caption_text)
                if not _INHERITED_TEMPLATE_MODE:
                    run.font.size = Pt(10.5)
                set_run_font_simsun(run)
                caption.paragraph_format.space_after = Pt(8)
                return True
            except Exception:
                return False

        def add_planned_blocks(chapter_id: str) -> None:
            """将已规划的表格和生成图表插入到对应章节位置。"""
            planned = blocks_by_chapter.get(str(chapter_id), [])
            if not planned:
                return
            for index, block in enumerate(planned, start=1):
                block_name = block.get("block_name") or block.get("name") or block.get("asset_name") or "文档块"
                block_type = block.get("block_type") or block.get("type") or "block"
                placeholder = block.get("placeholder") or block.get("fallback_placeholder") or ""
                asset = generated_asset_for_block(block, str(chapter_id), str(block_name))
                if block_type in visual_block_types:
                    if add_generated_visual_asset(asset, str(block_name), str(chapter_id), index):
                        continue
                    if placeholder:
                        add_markdown_paragraph(str(placeholder))

                table_schema = block.get("table_schema") or {}
                columns = table_schema.get("columns") if isinstance(table_schema, dict) else []
                if columns:
                    title_p = doc.add_paragraph()
                    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    add_markdown_runs(title_p, str(block_name))
                    add_three_line_table([
                        [str(column) for column in columns],
                        ["〖待补充〗"] * len(columns),
                    ])

                chart_schema = block.get("chart_schema") or {}
                if block_type in visual_block_types and isinstance(chart_schema, dict) and (chart_schema.get("nodes") or chart_schema.get("edges")) and not asset:
                    add_markdown_paragraph(f"〖插入图表：{block_name}〗")

                commitment_schema = block.get("commitment_schema") or {}
                if isinstance(commitment_schema, dict) and commitment_schema.get("items"):
                    add_markdown_paragraph(f"承诺书：{block_name}")
                    for item in commitment_schema.get("items", []):
                        add_markdown_paragraph(f"- {item}")

        # 递归构建文档内容（章节和内容）
        def add_outline_items(items, level: int = 1):
            for item in items:
                # 章节标题
                if level <= 3:
                    heading = doc.add_heading(f"{item.id} {item.title}", level=level)
                    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    if not _INHERITED_TEMPLATE_MODE:
                        for hr in heading.runs:
                            hr.font.name = "宋体"
                            rr = hr._element.rPr
                            if rr is not None and rr.rFonts is not None:
                                rr.rFonts.set(qn("w:eastAsia"), "宋体")
                else:
                    para = doc.add_paragraph()
                    run = para.add_run(f"{item.id} {item.title}")
                    if not _INHERITED_TEMPLATE_MODE:
                        run.bold = True
                        run.font.name = "宋体"
                        rr = run._element.rPr
                        if rr is not None and rr.rFonts is not None:
                            rr.rFonts.set(qn("w:eastAsia"), "宋体")
                    para.paragraph_format.space_before = Pt(6)
                    para.paragraph_format.space_after = Pt(3)

                # 叶子节点内容
                if not item.children:
                    content = item.content or ""
                    if add_history_word_blocks_if_possible(item):
                        pass
                    elif content.strip():
                        add_markdown_content(content, item=item)
                    add_planned_blocks(item.id)
                else:
                    add_planned_blocks(item.id)
                    add_outline_items(item.children, level + 1)

        add_outline_items(request.outline)

        # 输出到内存并返回
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        filename = f"{request.project_name or '标书文档'}.docx"
        if request.export_dir:
            export_dir = Path(request.export_dir).expanduser()
            export_dir.mkdir(parents=True, exist_ok=True)
            if not export_dir.is_dir():
                raise RuntimeError(f"保存路径不是目录：{export_dir}")

            safe_filename = sanitize_docx_filename(filename)
            output_path = export_dir / safe_filename
            output_path.write_bytes(buffer.getvalue())
            return JSONResponse({
                "success": True,
                "message": "Word 文件已保存到指定目录",
                "file_path": str(output_path),
                "filename": safe_filename,
            })

        # 使用 RFC 5987 格式对文件名进行 URL 编码，避免非 ASCII 字符导致的编码错误
        encoded_filename = quote(filename)
        content_disposition = f"attachment; filename*=UTF-8''{encoded_filename}"
        headers = {
            "Content-Disposition": content_disposition
        }

        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers=headers
        )
    except HTTPException:
        raise
    except Exception as e:
        # 打印详细错误信息到控制台，方便排查
        import traceback
        print("导出Word失败:", str(e))
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"导出Word失败: {str(e)}")
    finally:
        _INHERITED_TEMPLATE_MODE = previous_template_mode
