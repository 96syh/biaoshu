import importlib
import asyncio
import json
import os
import sqlite3
import sys
import tempfile
import unittest
from pathlib import Path


class BackendContractImportTests(unittest.TestCase):
    def test_schema_reexports_stay_compatible(self):
        from backend.app.models import schemas
        from backend.app.models.schema_defs.analysis import AnalysisReport as SplitAnalysisReport

        self.assertIs(schemas.AnalysisReport, SplitAnalysisReport)
        self.assertEqual(schemas.ConfigRequest().provider, "litellm")
        self.assertEqual(schemas.OutlineItem(id="1", title="章", description="").id, "1")

    def test_prompt_manager_reexports_stay_compatible(self):
        from backend.app.utils import prompt_manager

        self.assertIsInstance(prompt_manager.get_analysis_report_schema(), dict)
        self.assertIsInstance(prompt_manager.get_reference_bid_style_profile_schema(), dict)
        system_prompt, user_prompt = prompt_manager.generate_outline_prompt("项目概述", "评分要求")
        self.assertIn("目录", system_prompt)
        self.assertIn("项目概述", user_prompt)

    def test_openai_service_facade_keeps_public_methods(self):
        from backend.app.services.openai_service import OpenAIService

        service = OpenAIService()
        for name in (
            "generate_analysis_report",
            "generate_reference_bid_style_profile",
            "generate_document_blocks_plan",
            "generate_consistency_revision_report",
            "generate_response_matrix",
            "generate_compliance_review",
            "generate_content_for_outline",
            "_generate_chapter_content",
            "generate_outline_v2",
            "stream_chat_completion",
            "verify_current_endpoint",
        ):
            self.assertTrue(hasattr(service, name), name)

    def test_default_app_routes_exclude_optional_search_and_legacy_expand(self):
        os.environ["ENABLE_SEARCH_ROUTER"] = "0"
        os.environ["ENABLE_LEGACY_EXPAND_ROUTER"] = "0"
        sys.modules.pop("backend.app.main", None)

        main = importlib.import_module("backend.app.main")
        paths = {route.path for route in main.app.routes}

        self.assertIn("/api/document/upload", paths)
        self.assertIn("/api/document/upload-text", paths)
        self.assertIn("/api/document/source-preview/{source_preview_id}", paths)
        self.assertIn("/api/outline/generate", paths)
        self.assertIn("/api/projects", paths)
        self.assertFalse(any(path.startswith("/api/search") for path in paths))
        self.assertFalse(any(path.startswith("/api/expand") for path in paths))
        self.assertNotIn("duckduckgo_search", sys.modules)

    def test_runtime_data_defaults_to_artifacts(self):
        from backend.app.services.file_service import FileService
        from backend.app.services.generation_cache_service import GenerationCacheService
        from backend.app.services.history_case_service import HistoryCaseService
        from backend.app.services.project_service import ProjectService

        self.assertIn("artifacts/data/generated_assets", str(FileService.GENERATED_ASSET_DIR))
        self.assertIn("artifacts/data/generation_cache", str(GenerationCacheService.CACHE_DIR))
        self.assertIn("artifacts/data/history_cases/pageindex_trees", str(HistoryCaseService.PAGEINDEX_TREE_ROOT))
        self.assertIn("artifacts/runtime/projects.json", str(ProjectService.STORE_PATH))

    def test_model_gateway_concurrency_env_has_safe_floor(self):
        from backend.app.services.model_gateway_service import ModelGatewayService

        original_value = os.environ.get("YIBIAO_MODEL_CONCURRENCY")
        try:
            os.environ["YIBIAO_MODEL_CONCURRENCY"] = "0"
            self.assertEqual(ModelGatewayService._model_concurrency_limit(), 1)
            os.environ["YIBIAO_MODEL_CONCURRENCY"] = "3"
            self.assertEqual(ModelGatewayService._model_concurrency_limit(), 3)
        finally:
            if original_value is None:
                os.environ.pop("YIBIAO_MODEL_CONCURRENCY", None)
            else:
                os.environ["YIBIAO_MODEL_CONCURRENCY"] = original_value

    def test_generation_cache_round_trip(self):
        from backend.app.services.generation_cache_service import GenerationCacheService

        original_dir = GenerationCacheService.CACHE_DIR
        original_enabled = os.environ.get("YIBIAO_ENABLE_GENERATION_CACHE")
        with tempfile.TemporaryDirectory() as temp_dir:
            GenerationCacheService.CACHE_DIR = Path(temp_dir)
            os.environ["YIBIAO_ENABLE_GENERATION_CACHE"] = "1"
            try:
                key = GenerationCacheService.build_key("analysis", "model-a", {"x": 1})
                self.assertIsNone(GenerationCacheService.get("analysis", key))
                GenerationCacheService.set("analysis", key, {"ok": True})
                self.assertEqual(GenerationCacheService.get("analysis", key), {"ok": True})
            finally:
                GenerationCacheService.CACHE_DIR = original_dir
                if original_enabled is None:
                    os.environ.pop("YIBIAO_ENABLE_GENERATION_CACHE", None)
                else:
                    os.environ["YIBIAO_ENABLE_GENERATION_CACHE"] = original_enabled

    def test_scheme_outline_items_include_technical_composition_children(self):
        from backend.app.services.fallback_generation import FallbackGenerationMixin

        report = {
            "bid_document_requirements": {
                "composition": [
                    {
                        "id": "BD-TECH",
                        "title": "技术标",
                        "volume_id": "V-TECH",
                        "chapter_type": "technical",
                        "source_ref": "BD-SRC-01",
                        "children": [
                            {"id": "BD-TECH-01", "order": 1, "title": "服务范围"},
                            {"id": "BD-TECH-02", "order": 2, "title": "服务内容"},
                            {"id": "BD-TECH-03", "order": 3, "title": "质量承诺及措施"},
                        ],
                    }
                ],
                "selected_generation_target": {
                    "target_id": "BD-TECH",
                    "target_title": "技术标",
                    "parent_composition_id": "BD-TECH",
                    "base_outline_items": [],
                },
                "scheme_or_technical_outline_requirements": [],
            }
        }

        titles = [item["title"] for item in FallbackGenerationMixin._collect_scheme_outline_items(report)]

        self.assertEqual(titles[:3], ["服务范围", "服务内容", "质量承诺及措施"])

    def test_outline_guard_uses_tender_required_technical_chapters(self):
        from backend.app.services.generation.outline import OutlineGenerationMixin

        report = {
            "bid_document_requirements": {
                "composition": [
                    {
                        "id": "BD-TECH",
                        "title": "技术标",
                        "volume_id": "V-TECH",
                        "chapter_type": "technical",
                        "children": [
                            {"id": "BD-TECH-01", "order": 1, "title": "服务范围"},
                            {"id": "BD-TECH-02", "order": 2, "title": "服务内容"},
                        ],
                    }
                ],
                "selected_generation_target": {
                    "target_id": "BD-TECH",
                    "target_title": "技术标",
                    "parent_composition_id": "BD-TECH",
                    "base_outline_items": [],
                },
                "scheme_or_technical_outline_requirements": [],
            },
            "response_matrix": {"items": []},
        }

        tech_only_outline, changed, _ = OutlineGenerationMixin._apply_scheme_outline_guard(
            [{"id": "1", "title": "通用实施方案", "children": []}],
            report,
            "technical_only",
        )
        self.assertTrue(changed)
        self.assertEqual([item["title"] for item in tech_only_outline], ["服务范围", "服务内容"])

        full_bid_outline, changed, _ = OutlineGenerationMixin._apply_scheme_outline_guard(
            [{"id": "1", "title": "商务标"}, {"id": "2", "title": "技术标", "volume_id": "V-TECH"}],
            report,
            "full_bid",
        )
        self.assertTrue(changed)
        self.assertEqual([item["title"] for item in full_bid_outline[1]["children"]], ["服务范围", "服务内容"])

    def test_docx_html_preview_defaults_to_enabled_for_source_location(self):
        from backend.app.services.file_service import FileService

        original_value = os.environ.pop("YIBIAO_ENABLE_DOCX_HTML_PREVIEW", None)
        try:
            self.assertTrue(FileService.docx_html_preview_enabled())
            os.environ["YIBIAO_ENABLE_DOCX_HTML_PREVIEW"] = "0"
            self.assertFalse(FileService.docx_html_preview_enabled())
        finally:
            if original_value is not None:
                os.environ["YIBIAO_ENABLE_DOCX_HTML_PREVIEW"] = original_value
            else:
                os.environ.pop("YIBIAO_ENABLE_DOCX_HTML_PREVIEW", None)

    def test_source_preview_saved_upload_path_stays_under_upload_dir(self):
        from backend.app.config import settings
        from backend.app.services.file_service import FileService

        original_upload_dir = settings.upload_dir
        with tempfile.TemporaryDirectory() as temp_dir:
            settings.upload_dir = temp_dir
            safe_path = Path(temp_dir) / "sample.docx"
            safe_path.write_text("placeholder", encoding="utf-8")
            try:
                self.assertEqual(FileService._resolve_saved_upload_path("sample.docx"), str(safe_path.resolve()))
                with self.assertRaises(Exception):
                    FileService._resolve_saved_upload_path("../sample.docx")
            finally:
                settings.upload_dir = original_upload_dir

    def test_empty_docx_preview_html_is_not_treated_as_renderable(self):
        from backend.app.services.file_service import FileService

        self.assertFalse(FileService.source_preview_html_has_text('<div class="docx-source-preview"><p>&nbsp;</p></div>'))
        self.assertTrue(FileService.source_preview_html_has_text('<div class="docx-source-preview"><p>服务质量保证</p></div>'))

    def test_docx_line_spacing_conversion_does_not_emit_raw_emu_values(self):
        from backend.app.services.file_service import FileService

        class EmuLength:
            pt = 22

            def __float__(self):
                return 279400.0

        self.assertEqual(FileService._line_spacing_to_css(EmuLength()), "22.0pt")
        self.assertEqual(FileService._line_spacing_to_css(1.5), "1.50")
        self.assertEqual(FileService._line_spacing_to_css(279400), "")

    def test_history_case_legacy_artifact_paths_resolve_to_artifacts(self):
        from backend.app.services.history_case_service import HistoryCaseService

        original_artifact_root = HistoryCaseService.HISTORY_ARTIFACT_ROOT
        original_legacy_root = HistoryCaseService.LEGACY_HISTORY_ARTIFACT_ROOT
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_root = Path(temp_dir)
            HistoryCaseService.HISTORY_ARTIFACT_ROOT = temp_root / "artifacts" / "data" / "history_cases"
            HistoryCaseService.LEGACY_HISTORY_ARTIFACT_ROOT = temp_root / "backend" / "data" / "history_cases"
            target = HistoryCaseService.HISTORY_ARTIFACT_ROOT / "markdown" / "case.md"
            target.parent.mkdir(parents=True)
            target.write_text("history markdown", encoding="utf-8")
            legacy_path = HistoryCaseService.LEGACY_HISTORY_ARTIFACT_ROOT / "markdown" / "case.md"

            try:
                self.assertEqual(HistoryCaseService.load_markdown(str(legacy_path)), "history markdown")
                self.assertEqual(
                    HistoryCaseService._normalize_artifact_path_for_response(str(legacy_path)),
                    str(target),
                )
            finally:
                HistoryCaseService.HISTORY_ARTIFACT_ROOT = original_artifact_root
                HistoryCaseService.LEGACY_HISTORY_ARTIFACT_ROOT = original_legacy_root

    def test_history_chapter_reference_extracts_matching_section(self):
        from backend.app.services.history_case_service import HistoryCaseService

        markdown = (
            "# 技术标\n\n"
            "## 服务范围\n旧服务范围正文\n\n"
            "## 质量保证措施\n质量目标明确，建立校审、复核和问题闭环机制。\n\n"
            "## 进度保证措施\n进度计划正文"
        )

        section = HistoryCaseService._extract_markdown_section(markdown, "质量保证措施")

        self.assertIn("质量目标明确", section)
        self.assertNotIn("进度计划正文", section)

    def test_history_chapter_reference_preserves_images_and_tables(self):
        from backend.app.services.history_case_service import HistoryCaseService

        markdown = (
            "# 技术标\n\n"
            "## 第一章 服务质量保证\n"
            "![组织架构图](assets/org.png)\n\n"
            "| 序号 | 措施 |\n"
            "| --- | --- |\n"
            "| 1 | 校审复核 |\n"
        )

        section = HistoryCaseService._extract_markdown_section(markdown, "服务质量保证")
        inventory = HistoryCaseService._reference_block_inventory(section)

        self.assertIn("![组织架构图](assets/org.png)", section)
        self.assertIn("| 序号 | 措施 |", section)
        self.assertTrue(inventory["has_image"])
        self.assertTrue(inventory["has_table"])

    def test_markdown_fallback_reference_keeps_only_text_before_content_generation(self):
        from backend.app.services.history_case_service import HistoryCaseService

        reference = (
            "质量保证措施如下：\n\n"
            "![组织架构图](assets/org.png)\n\n"
            "| 序号 | 措施 |\n"
            "| --- | --- |\n"
            "| 1 | 校审复核 |\n\n"
            "<figure><img src=\"assets/process.png\" /></figure>\n\n"
            "<table><tr><td>资料</td></tr></table>\n\n"
            "项目组将建立问题闭环机制。\n"
        )

        stripped = HistoryCaseService._strip_non_text_markdown_from_reference(reference)

        self.assertIn("质量保证措施如下", stripped)
        self.assertIn("项目组将建立问题闭环机制", stripped)
        self.assertNotIn("![组织架构图](assets/org.png)", stripped)
        self.assertNotIn("| 序号 | 措施 |", stripped)
        self.assertNotIn("| 1 | 校审复核 |", stripped)
        self.assertNotIn("<figure>", stripped)
        self.assertNotIn("<table>", stripped)

    def test_chapter_content_prompt_carries_history_rewrite_rules(self):
        from backend.app.utils import prompt_manager

        system_prompt, user_prompt = prompt_manager.generate_chapter_content_prompt(
            chapter={"id": "c1", "title": "质量保证措施", "description": "响应质量评分项"},
            parent_chapters=[],
            sibling_chapters=[],
            project_overview="油库设计服务",
            analysis_report={"project": {"name": "当前项目"}},
            history_reference_drafts=[
                {
                    "match_level": "high",
                    "project_title": "历史项目",
                    "reference_text": "历史质量保证正文",
                }
            ],
        )

        self.assertIn("历史正文主稿规则", system_prompt)
        self.assertIn("当前招标文件", system_prompt)
        self.assertIn("不得扩写超过历史 reference_text 字数的 130%", system_prompt)
        self.assertIn("历史企业固定资料复用", system_prompt)
        self.assertIn("人员姓名、职称、证书编号", system_prompt)
        self.assertIn("图片语法", system_prompt)
        self.assertIn("表格", system_prompt)
        self.assertIn("history_reference_drafts", user_prompt)
        self.assertIn("历史质量保证正文", user_prompt)

    def test_chapter_patch_prompt_emits_patch_contract(self):
        from backend.app.utils import prompt_manager

        system_prompt, user_prompt = prompt_manager.generate_chapter_patch_prompt(
            chapter={"id": "1", "title": "服务范围"},
            parent_chapters=[],
            project_overview="油库设计服务",
            analysis_report={"technical_scoring_items": [{"name": "服务范围", "standard": "响应全面"}]},
            response_matrix={"items": []},
            history_reference_draft={"matched_blocks": [{"id": "b-1", "text": "历史服务范围"}]},
        )

        self.assertIn("补丁指令", system_prompt)
        self.assertIn("matched_blocks", user_prompt)
        self.assertIn("当前目录标题", system_prompt)
        self.assertIn("不要替换成 〖待补充〗", system_prompt)

    def test_personnel_chapter_reference_terms_include_fixed_material_aliases(self):
        from backend.app.services.history_case_service import HistoryCaseService

        terms = HistoryCaseService._chapter_reference_terms(
            "5.1 项目负责人",
            "拟投入的服务人员、项目负责人、人员证书和职称要求",
        )

        self.assertIn("项目负责人", terms)
        self.assertIn("拟投入人员", terms)
        self.assertIn("拟委任的主要人员", terms)
        self.assertIn("主要人员汇总表", terms)

    def test_personnel_semantic_tokens_match_history_headings(self):
        from backend.app.services.history_case_service import HistoryCaseService

        title_tokens = HistoryCaseService._chapter_reference_semantic_tokens("项目组成人员详情")
        heading_tokens = HistoryCaseService._chapter_reference_semantic_tokens("拟委任的主要人员汇总表")

        self.assertTrue(title_tokens & heading_tokens)

    def test_history_patch_operations_apply_text_replacements(self):
        from backend.app.services.openai_service import OpenAIService

        operations = [
            {"op": "replace_text", "from": "历史项目", "to": "当前项目"},
            {"op": "append_text", "text": "补充评分项响应。"},
        ]

        patched = OpenAIService._apply_history_patch_operations("历史项目服务范围。", operations)

        self.assertIn("当前项目服务范围。", patched)
        self.assertIn("补充评分项响应。", patched)

    def test_history_patch_operations_apply_to_blocks_by_id(self):
        from backend.app.services.openai_service import OpenAIService

        blocks = [
            {"id": "p-1", "type": "paragraph", "text": "历史项目服务范围", "markdown": "历史项目服务范围", "html": '<p data-history-block-id="p-1">历史项目服务范围</p>'},
            {"id": "img-1", "type": "paragraph", "text": "", "markdown": "![组织架构](assets/org.png)", "html": '<figure data-history-block-id="img-1"><img src="assets/org.png" /></figure>'},
            {"id": "p-2", "type": "paragraph", "text": "后续说明", "markdown": "后续说明", "html": '<p data-history-block-id="p-2">后续说明</p>'},
        ]
        operations = [
            {"op": "replace_text", "block_id": "p-1", "from": "历史项目", "to": "当前项目"},
            {"op": "insert_after", "after_block_id": "p-1", "text": "补充评分项响应。"},
            {"op": "update_caption", "block_id": "img-1", "caption": "本项目服务组织架构图"},
            {"op": "move_block_after", "block_id": "img-1", "after_block_id": "p-2"},
        ]

        patched = OpenAIService._apply_history_patch_to_blocks(blocks, operations)
        markdown = "\n".join(str(block.get("markdown") or "") for block in patched)

        self.assertIn("当前项目服务范围", markdown)
        self.assertIn("补充评分项响应。", markdown)
        self.assertIn("本项目服务组织架构图", markdown)
        self.assertEqual([block["id"] for block in patched][-1], "img-1")

    def test_history_patch_inserted_markdown_table_renders_as_html_table(self):
        from backend.app.services.openai_service import OpenAIService

        blocks = [
            {"id": "p-1", "type": "paragraph", "text": "进度目标", "markdown": "进度目标", "html": '<p data-history-block-id="p-1">进度目标</p>'},
        ]
        operations = [{
            "op": "insert_after",
            "after_block_id": "p-1",
            "text": "进度目标响应表 | 序号 | 进度事项 | 招标文件要求 | 本章响应目标 | 管控方式 | | 1 | 服务期限 | 至2026年12月31日 | 动态组织设计服务 | 建立任务台账 |",
        }]

        patched = OpenAIService._apply_history_patch_to_blocks(blocks, operations)
        html = patched[1]["html"]

        self.assertIn("进度目标响应表", html)
        self.assertIn("<table", html)
        self.assertIn("<th>序号</th>", html)
        self.assertIn("<td>服务期限</td>", html)
        self.assertNotIn("| 序号 |", html)

    def test_history_patch_output_drops_all_section_headings(self):
        from backend.app.services.openai_service import OpenAIService

        blocks = [
            {"id": "h-1", "type": "heading", "text": "总体实施路径与资源协同安排", "markdown": "## 总体实施路径与资源协同安排", "html": '<h2 data-history-block-id="h-1">总体实施路径与资源协同安排</h2>'},
            {"id": "p-1", "type": "paragraph", "text": "本节围绕资源配置展开。", "markdown": "本节围绕资源配置展开。", "html": '<p data-history-block-id="p-1">本节围绕资源配置展开。</p>'},
            {"id": "h-2", "type": "heading", "text": "设计服务流程图", "markdown": "### 设计服务流程图", "html": '<h3 data-history-block-id="h-2">设计服务流程图</h3>'},
            {"id": "p-2", "type": "paragraph", "text": "流程说明。", "markdown": "流程说明。", "html": '<p data-history-block-id="p-2">流程说明。</p>'},
        ]

        stripped = OpenAIService._strip_history_heading_blocks_from_content(blocks)
        markdown = "\n".join(str(block.get("markdown") or "") for block in stripped)

        self.assertEqual([block["id"] for block in stripped], ["p-1", "p-2"])
        self.assertNotIn("总体实施路径与资源协同安排", markdown)
        self.assertNotIn("设计服务流程图", markdown)

        compact = OpenAIService._compact_history_reference_draft_for_model({"matched_blocks": blocks})
        self.assertNotIn("总体实施路径与资源协同安排", compact["reference_text"])
        self.assertNotIn("设计服务流程图", compact["reference_text"])
        self.assertEqual([block["id"] for block in compact["preserved_word_blocks"]], ["p-1", "p-2"])

    def test_generated_markdown_heading_lines_are_removed(self):
        from backend.app.services.openai_service import OpenAIService

        content = "## 设计服务流程图\n本节说明流程。\n\n### 专业协同与接口管理\n协同内容。"

        stripped = OpenAIService._strip_generated_markdown_headings(content)

        self.assertNotIn("## 设计服务流程图", stripped)
        self.assertNotIn("### 专业协同与接口管理", stripped)
        self.assertIn("本节说明流程。", stripped)
        self.assertIn("协同内容。", stripped)

    def test_primary_history_draft_requires_word_heading_blocks(self):
        from backend.app.services.openai_service import OpenAIService

        selected = OpenAIService._select_primary_history_draft([
            {"match_level": "high", "matched_blocks": [], "reference_source": "markdown_fallback"},
            {"match_level": "medium", "matched_blocks": [{"id": "b-1"}], "reference_source": "word_heading_blocks"},
        ])

        self.assertEqual(selected["matched_blocks"][0]["id"], "b-1")
        self.assertIsNone(OpenAIService._select_primary_history_draft([
            {"match_level": "high", "matched_blocks": [], "reference_source": "markdown_fallback"},
        ]))

    def test_history_reference_prompt_payload_omits_heavy_word_blocks(self):
        from backend.app.services.openai_service import OpenAIService

        draft = {
            "match_level": "high",
            "score": 0.8,
            "reference_source": "word_heading_blocks",
            "matched_blocks": [
                {
                    "id": "b-1",
                    "type": "paragraph",
                    "text": "历史项目服务范围",
                    "markdown": "历史项目服务范围",
                    "html": '<p data-history-block-id="b-1">历史项目服务范围</p>',
                    "docx_xml": "<w:p>very-heavy-word-xml</w:p>",
                },
                {
                    "id": "b-2",
                    "type": "table",
                    "text": "序号 | 姓名 | 证书编号",
                    "markdown": "| 序号 | 姓名 | 证书编号 |\n| --- | --- | --- |\n| 1 | 张三 | A001 |",
                    "html": "<table>very-heavy-table-html</table>",
                    "docx_xml": "<w:tbl>very-heavy-table-xml</w:tbl>",
                    "rows": [["序号", "姓名", "证书编号"], ["1", "张三", "A001"]],
                },
                {
                    "id": "b-3",
                    "type": "paragraph",
                    "text": "",
                    "markdown": "![组织架构](assets/org.png)",
                    "html": '<figure><img src="data:image/png;base64,abcdef" /></figure>',
                    "asset_ids": ["img-1"],
                },
            ],
        }

        compact = OpenAIService._compact_history_reference_draft_for_model(draft)
        payload = json.dumps(compact, ensure_ascii=False)

        self.assertIn("历史项目服务范围", compact["reference_text"])
        self.assertIn("Word表格已从prompt省略", payload)
        self.assertIn("Word图片已从prompt省略", payload)
        self.assertNotIn("very-heavy-word-xml", payload)
        self.assertNotIn("very-heavy-table-html", payload)
        self.assertNotIn("base64,abcdef", payload)
        self.assertNotIn("| 序号 | 姓名 |", compact["reference_text"])

    def test_history_text_reference_excludes_table_blocks_from_plain_content(self):
        from backend.app.services.openai_service import OpenAIService
        from backend.app.services.history_case_service import HistoryCaseService

        blocks = [
            {
                "id": "p-1",
                "type": "paragraph",
                "text": "包括工艺、结构、建筑、管道、设备、电气、仪表、电信、给排水、暖通等专业设计及现场技术服务。",
                "markdown": "包括工艺、结构、建筑、管道、设备、电气、仪表、电信、给排水、暖通等专业设计及现场技术服务。",
                "html": "<p>包括工艺、结构、建筑、管道、设备、电气、仪表、电信、给排水、暖通等专业设计及现场技术服务。</p>",
            },
            {
                "id": "t-1",
                "type": "table",
                "text": "服务阶段 主要工作 输入资料 输出成果 配合对象 注意事项 任务接收与启动 接收招标人委托任务",
                "markdown": "服务阶段 主要工作 输入资料 输出成果 配合对象 注意事项 任务接收与启动 接收招标人委托任务",
                "html": "<table><tr><th>服务阶段</th><th>主要工作</th></tr><tr><td>任务接收与启动</td><td>接收招标人委托任务</td></tr></table>",
            },
        ]

        plain_content = OpenAIService._history_text_reference_from_blocks(blocks)
        html_content = HistoryCaseService._blocks_to_html(blocks)

        self.assertIn("包括工艺", plain_content)
        self.assertNotIn("服务阶段 主要工作 输入资料", plain_content)
        self.assertIn("<table", html_content)
        self.assertIn("任务接收与启动", html_content)

    def test_history_word_blocks_reuse_media_except_blind_bid_sensitive_blocks(self):
        from backend.app.services.openai_service import OpenAIService

        blocks = [
            {"id": "b-1", "type": "paragraph", "text": "服务说明", "markdown": "服务说明", "html": '<p data-history-block-id="b-1">服务说明</p>'},
            {"id": "b-2", "type": "table", "text": "序号 | 姓名 | 联系方式", "markdown": "| 序号 | 姓名 | 联系方式 |", "html": '<table data-history-block-id="b-2"></table>'},
            {"id": "b-3", "type": "paragraph", "text": "", "markdown": "![组织架构](assets/org.png)", "html": '<figure><img src="assets/org.png" /></figure>', "asset_ids": ["img-1"]},
        ]

        reused = OpenAIService._prepare_history_blocks_for_reuse(blocks, chapter={"title": "组织机构"})
        blind_reused = OpenAIService._prepare_history_blocks_for_reuse(
            blocks,
            chapter={"title": "暗标组织机构", "description": "不得出现人员姓名、联系方式、Logo"},
        )

        self.assertEqual([block["id"] for block in reused], ["b-1", "b-2", "b-3"])
        self.assertIn("待人工复核", blind_reused[1]["markdown"])
        self.assertIn("待人工复核", blind_reused[2]["markdown"])
        self.assertNotIn("assets/org.png", "\n".join(str(block.get("markdown") or "") for block in blind_reused))

    def test_history_blocks_convert_to_markdown_and_html(self):
        from backend.app.services.history_case_service import HistoryCaseService

        blocks = [
            {"id": "b-1", "type": "heading", "text": "服务范围", "markdown": "## 服务范围", "html": "<h2 data-history-block-id=\"b-1\">服务范围</h2>"},
            {"id": "b-2", "type": "table", "text": "序号 | 内容", "markdown": "| 序号 | 内容 |\n| --- | --- |\n| 1 | 响应 |", "html": "<table data-history-block-id=\"b-2\"></table>"},
        ]

        self.assertIn("| 序号 | 内容 |", HistoryCaseService._blocks_to_markdown(blocks))
        self.assertIn("data-history-block-id", HistoryCaseService._blocks_to_html(blocks))

    def test_history_matching_anchors_only_heading_and_excludes_service_tables(self):
        from backend.app.services.history_case_service import HistoryCaseService

        blocks = [
            {
                "id": "b-1",
                "type": "heading",
                "level": 1,
                "text": "业绩文件",
                "markdown": "# 业绩文件",
                "html": '<h1 data-history-block-id="b-1">业绩文件</h1>',
                "heading_path": [{"id": "b-1", "level": 1, "title": "业绩文件"}],
            },
            {
                "id": "b-2",
                "type": "table",
                "level": 0,
                "text": "项目法人 | 服务范围 | 服务人员人数",
                "markdown": "| 项目法人 | 服务范围 | 服务人员人数 |",
                "html": '<table data-history-block-id="b-2"></table>',
                "heading_path": [{"id": "b-1", "level": 1, "title": "业绩文件"}],
            },
            {
                "id": "b-3",
                "type": "heading",
                "level": 1,
                "text": "服务范围、服务内容",
                "markdown": "# 服务范围、服务内容",
                "html": '<h1 data-history-block-id="b-3">服务范围、服务内容</h1>',
                "heading_path": [{"id": "b-3", "level": 1, "title": "服务范围、服务内容"}],
            },
            {
                "id": "b-4",
                "type": "paragraph",
                "level": 0,
                "text": "本节说明服务范围。",
                "markdown": "本节说明服务范围。",
                "html": '<p data-history-block-id="b-4">本节说明服务范围。</p>',
                "heading_path": [{"id": "b-3", "level": 1, "title": "服务范围、服务内容"}],
            },
            {
                "id": "b-5",
                "type": "heading",
                "level": 1,
                "text": "拟投入人员",
                "markdown": "# 拟投入人员",
                "html": '<h1 data-history-block-id="b-5">拟投入人员</h1>',
                "heading_path": [{"id": "b-5", "level": 1, "title": "拟投入人员"}],
            },
            {
                "id": "b-6",
                "type": "table",
                "level": 0,
                "text": "层级 | 职务 | 姓名",
                "markdown": "| 层级 | 职务 | 姓名 |",
                "html": '<table data-history-block-id="b-6"></table>',
                "heading_path": [{"id": "b-5", "level": 1, "title": "拟投入人员"}],
            },
        ]

        matched = HistoryCaseService._extract_matching_blocks(blocks, "服务范围、服务内容")

        self.assertEqual([block["id"] for block in matched], ["b-3", "b-4"])
        matched_text = HistoryCaseService._blocks_to_markdown(matched)
        self.assertIn("本节说明服务范围", matched_text)
        self.assertNotIn("项目法人", matched_text)
        self.assertNotIn("层级 | 职务 | 姓名", matched_text)

    def test_history_matching_uses_heading_path_for_descendants(self):
        from backend.app.services.history_case_service import HistoryCaseService

        root_path = [{"id": "b-1", "level": 1, "title": "服务范围、服务内容"}]
        blocks = [
            {"id": "b-1", "type": "heading", "level": 1, "text": "服务范围、服务内容", "markdown": "# 服务范围、服务内容", "heading_path": root_path},
            {
                "id": "b-2",
                "type": "heading",
                "level": 2,
                "text": "服务范围",
                "markdown": "## 服务范围",
                "heading_path": [*root_path, {"id": "b-2", "level": 2, "title": "服务范围"}],
            },
            {
                "id": "b-3",
                "type": "paragraph",
                "text": "历史服务范围正文",
                "markdown": "历史服务范围正文",
                "heading_path": [*root_path, {"id": "b-2", "level": 2, "title": "服务范围"}],
            },
            {
                "id": "b-4",
                "type": "heading",
                "level": 2,
                "text": "服务内容",
                "markdown": "## 服务内容",
                "heading_path": [*root_path, {"id": "b-4", "level": 2, "title": "服务内容"}],
            },
            {
                "id": "b-5",
                "type": "paragraph",
                "text": "历史服务内容正文",
                "markdown": "历史服务内容正文",
                "heading_path": [*root_path, {"id": "b-4", "level": 2, "title": "服务内容"}],
            },
            {
                "id": "b-6",
                "type": "heading",
                "level": 1,
                "text": "拟投入人员",
                "markdown": "# 拟投入人员",
                "heading_path": [{"id": "b-6", "level": 1, "title": "拟投入人员"}],
            },
        ]

        matched = HistoryCaseService._extract_matching_blocks(blocks, "服务范围、服务内容")

        self.assertEqual([block["id"] for block in matched], ["b-1", "b-2", "b-3", "b-4", "b-5"])

    def test_markdown_section_does_not_fallback_to_document_start(self):
        from backend.app.services.history_case_service import HistoryCaseService

        markdown = "# 技术文件\n\n# 业绩文件\n\n项目法人 | 服务范围 | 姓名 | 职称 | 证书"

        section = HistoryCaseService._extract_markdown_section(markdown, "服务范围")

        self.assertEqual(section, "")

    def test_docx_block_extraction_keeps_original_ooxml(self):
        import docx
        from backend.app.services.history_case_service import extract_docx_blocks

        with tempfile.TemporaryDirectory() as temp_dir:
            temp_root = Path(temp_dir)
            source_path = temp_root / "history.docx"
            document = docx.Document()
            document.add_heading("服务范围", level=1)
            document.add_paragraph("历史项目服务范围")
            table = document.add_table(rows=1, cols=2)
            table.rows[0].cells[0].text = "序号"
            table.rows[0].cells[1].text = "内容"
            document.save(source_path)

            payload = extract_docx_blocks(source_path, temp_root / "assets")
            blocks = payload["blocks"]

            self.assertTrue(any(block.get("docx_xml", "").startswith("<w:p") for block in blocks))
            self.assertTrue(any(block.get("docx_xml", "").startswith("<w:tbl") for block in blocks))
            self.assertTrue(all("heading_path" in block for block in blocks))
            self.assertEqual(blocks[1]["heading_path"][0]["title"], "服务范围")
            self.assertEqual(blocks[2]["heading_path"][0]["title"], "服务范围")

    def test_history_patch_updates_original_ooxml_text(self):
        from backend.app.services.openai_service import OpenAIService

        blocks = [
            {
                "id": "p-1",
                "type": "paragraph",
                "text": "历史项目服务范围",
                "markdown": "历史项目服务范围",
                "html": '<p data-history-block-id="p-1">历史项目服务范围</p>',
                "docx_xml": (
                    '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                    "<w:r><w:t>历史项目服务范围</w:t></w:r></w:p>"
                ),
            }
        ]

        patched = OpenAIService._apply_history_patch_to_blocks(
            blocks,
            [{"op": "replace_text", "block_id": "p-1", "from": "历史项目", "to": "当前项目"}],
        )

        self.assertIn("当前项目服务范围", patched[0]["docx_xml"])

    def test_word_export_uses_history_docx_as_style_template(self):
        import docx
        from docx.shared import Cm
        from backend.app.models.schemas import OutlineItem, WordExportRequest
        from backend.app.services.word_export_service import create_export_document

        with tempfile.TemporaryDirectory() as temp_dir:
            template_path = Path(temp_dir) / "template.docx"
            template = docx.Document()
            template.sections[0].left_margin = Cm(3.3)
            template.add_paragraph("旧正文")
            template.save(template_path)

            request = WordExportRequest(
                manual_review_confirmed=True,
                outline=[
                    OutlineItem(
                        id="1",
                        title="服务范围",
                        description="",
                        history_reference={"source_paths": {"source_docx_path": str(template_path)}},
                    )
                ],
            )

            doc, inherited_path = create_export_document(request)

            self.assertEqual(inherited_path, template_path)
            self.assertAlmostEqual(int(doc.sections[0].left_margin), int(Cm(3.3)), delta=200)
            self.assertNotIn("旧正文", "\n".join(paragraph.text for paragraph in doc.paragraphs))

    def test_word_export_reuses_history_header_footer_without_overwriting(self):
        import docx
        from docx.enum.section import WD_SECTION
        from backend.app.models.schemas import OutlineItem, WordExportRequest
        from backend.app.services.word_export_service import create_word_export_response

        with tempfile.TemporaryDirectory() as temp_dir:
            template_path = Path(temp_dir) / "template.docx"
            template = docx.Document()
            template.add_paragraph("封面正文")
            main_section = template.add_section(WD_SECTION.NEW_PAGE)
            main_section.header.is_linked_to_previous = False
            main_section.header.paragraphs[0].text = "历史页眉-技术文件"
            main_section.footer.is_linked_to_previous = False
            main_section.footer.paragraphs[0].text = "历史页脚-PAGE"
            template.add_paragraph("历史正文")
            template.save(template_path)

            request = WordExportRequest(
                project_name="测试导出",
                manual_review_confirmed=True,
                export_dir=temp_dir,
                outline=[
                    OutlineItem(
                        id="1",
                        title="服务范围",
                        description="",
                        content="本节说明服务范围。",
                        history_reference={"source_paths": {"source_docx_path": str(template_path)}},
                    )
                ],
            )

            response = asyncio.run(create_word_export_response(request))
            payload = json.loads(response.body.decode("utf-8"))
            exported = docx.Document(payload["file_path"])
            header_text = "\n".join(p.text for p in exported.sections[0].header.paragraphs)
            footer_text = "\n".join(p.text for p in exported.sections[0].footer.paragraphs)

            self.assertIn("测试导出-技术文件", header_text)
            self.assertNotIn("历史页眉-技术文件", header_text)
            self.assertIn("历史页脚-PAGE", footer_text)
            self.assertNotIn("第 ", footer_text)
            self.assertNotIn("共 ", footer_text)

    def test_word_export_template_toc_uses_history_like_depth_two(self):
        import docx
        import zipfile
        from backend.app.models.schemas import OutlineItem, WordExportRequest
        from backend.app.services.word_export_service import create_word_export_response

        with tempfile.TemporaryDirectory() as temp_dir:
            template_path = Path(temp_dir) / "template.docx"
            template = docx.Document()
            template.sections[0].header.paragraphs[0].text = "历史页眉"
            template.sections[0].footer.paragraphs[0].text = "历史页脚"
            template.add_paragraph("历史正文")
            template.save(template_path)

            request = WordExportRequest(
                project_name="测试导出",
                manual_review_confirmed=True,
                export_dir=temp_dir,
                outline=[
                    OutlineItem(
                        id="1",
                        title="服务范围、服务内容",
                        description="",
                        history_reference={"source_paths": {"source_docx_path": str(template_path)}},
                        children=[
                            OutlineItem(id="1.1", title="服务范围", description="", content="本节说明服务范围。")
                        ],
                    )
                ],
            )

            response = asyncio.run(create_word_export_response(request))
            payload = json.loads(response.body.decode("utf-8"))
            exported = docx.Document(payload["file_path"])
            heading_texts = [
                paragraph.text
                for paragraph in exported.paragraphs
                if getattr(paragraph.style, "name", "").startswith("Heading")
            ]
            with zipfile.ZipFile(payload["file_path"]) as package:
                document_xml = package.read("word/document.xml").decode("utf-8")

            self.assertIn("目  录", "\n".join(paragraph.text for paragraph in exported.paragraphs))
            self.assertIn('TOC \\o "1-2" \\h \\u', document_xml)
            self.assertNotIn("目  录", heading_texts)

    def test_word_export_updates_history_header_project_title(self):
        import docx
        import zipfile
        from backend.app.models.schemas import OutlineItem, WordExportRequest
        from backend.app.services.word_export_service import create_word_export_response

        with tempfile.TemporaryDirectory() as temp_dir:
            template_path = Path(temp_dir) / "template.docx"
            template = docx.Document()
            header = template.sections[0].header
            header.paragraphs[0].add_run("2025-2027 年工程设计服务框架项目投标文件-技术文件")
            pict = header.paragraphs[0]._p.add_r()
            pict.append(docx.oxml.parse_xml(
                '<w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
            ))
            template.add_paragraph("历史正文")
            template.save(template_path)

            request = WordExportRequest(
                project_name="中国石油辽宁销售公司2026年油库、加油站等工程设计服务项目",
                manual_review_confirmed=True,
                export_dir=temp_dir,
                outline=[
                    OutlineItem(
                        id="1",
                        title="服务范围",
                        description="",
                        content="本节说明服务范围。",
                        history_reference={"source_paths": {"source_docx_path": str(template_path)}},
                    )
                ],
            )

            response = asyncio.run(create_word_export_response(request))
            payload = json.loads(response.body.decode("utf-8"))
            exported = docx.Document(payload["file_path"])
            header_text = "\n".join(p.text for p in exported.sections[0].header.paragraphs)
            with zipfile.ZipFile(payload["file_path"]) as package:
                header_xml = "".join(
                    package.read(name).decode("utf-8")
                    for name in package.namelist()
                    if name.startswith("word/header")
                )

            self.assertIn("中国石油辽宁销售公司2026年油库、加油站等工程设计服务项目-技术文件", header_text)
            self.assertNotIn("2025-2027 年工程设计服务框架项目投标文件-技术文件", header_text)
            self.assertIn("<w:pict", header_xml)

    def test_word_export_prefers_generated_content_over_history_blocks(self):
        import docx
        from backend.app.models.schemas import OutlineItem, WordExportRequest
        from backend.app.services.word_export_service import create_word_export_response

        with tempfile.TemporaryDirectory() as temp_dir:
            template_path = Path(temp_dir) / "template.docx"
            template = docx.Document()
            template.add_paragraph("历史模板正文")
            template.save(template_path)

            block_json_path = Path(temp_dir) / "blocks.json"
            block_json_path.write_text(json.dumps({
                "blocks": [
                    {
                        "id": "b-1",
                        "type": "paragraph",
                        "text": "历史旧正文",
                        "markdown": "历史旧正文",
                        "docx_xml": (
                            '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                            "<w:r><w:t>历史旧正文</w:t></w:r></w:p>"
                        ),
                    }
                ]
            }), encoding="utf-8")

            request = WordExportRequest(
                project_name="测试导出",
                manual_review_confirmed=True,
                export_dir=temp_dir,
                outline=[
                    OutlineItem(
                        id="1",
                        title="服务范围",
                        description="",
                        content="模型生成正文应当被导出。",
                        history_reference={
                            "matched_block_ids": ["b-1"],
                            "source_paths": {
                                "source_docx_path": str(template_path),
                                "block_json_path": str(block_json_path),
                            },
                        },
                    )
                ],
            )

            response = asyncio.run(create_word_export_response(request))
            payload = json.loads(response.body.decode("utf-8"))
            exported = docx.Document(payload["file_path"])
            text = "\n".join(paragraph.text for paragraph in exported.paragraphs)

            self.assertIn("模型生成正文应当被导出。", text)
            self.assertNotIn("历史旧正文", text)

    def test_word_export_keeps_content_html_tables_with_full_borders(self):
        import docx
        import zipfile
        from backend.app.models.schemas import OutlineItem, WordExportRequest
        from backend.app.services.word_export_service import create_word_export_response

        with tempfile.TemporaryDirectory() as temp_dir:
            request = WordExportRequest(
                project_name="测试导出",
                manual_review_confirmed=True,
                export_dir=temp_dir,
                outline=[
                    OutlineItem(
                        id="5.2",
                        title="项目组成人员详情",
                        description="",
                        content="本节说明人员配置。",
                        content_html=(
                            "<table><tbody>"
                            "<tr><th>序号</th><th>本项目职务</th><th>姓名</th></tr>"
                            "<tr><td>1</td><td>项目经理</td><td>关义</td></tr>"
                            "</tbody></table>"
                        ),
                    )
                ],
            )

            response = asyncio.run(create_word_export_response(request))
            payload = json.loads(response.body.decode("utf-8"))
            exported = docx.Document(payload["file_path"])
            table_text = "\n".join(
                "\t".join(cell.text for cell in row.cells)
                for table in exported.tables
                for row in table.rows
            )
            with zipfile.ZipFile(payload["file_path"]) as package:
                document_xml = package.read("word/document.xml").decode("utf-8")

            self.assertIn("本节说明人员配置。", "\n".join(p.text for p in exported.paragraphs))
            self.assertIn("项目经理", table_text)
            self.assertIn("关义", table_text)
            self.assertIn("<w:insideH", document_xml)
            self.assertIn("<w:insideV", document_xml)
            self.assertNotIn('w:val="nil"', document_xml)

    def test_word_export_renumbers_outline_and_removes_duplicate_content_heading(self):
        import docx
        from backend.app.models.schemas import OutlineItem, WordExportRequest
        from backend.app.services.word_export_service import create_word_export_response

        with tempfile.TemporaryDirectory() as temp_dir:
            request = WordExportRequest(
                project_name="测试导出",
                manual_review_confirmed=True,
                export_dir=temp_dir,
                outline=[
                    OutlineItem(
                        id="4",
                        title="1 服务范围、服务内容",
                        description="",
                        children=[
                            OutlineItem(
                                id="4.1.1",
                                title="服务范围",
                                description="",
                                content="4.1.1服务范围\n本节说明服务范围。",
                            )
                        ],
                    )
                ],
            )

            response = asyncio.run(create_word_export_response(request))
            payload = json.loads(response.body.decode("utf-8"))
            exported = docx.Document(payload["file_path"])
            text = "\n".join(paragraph.text for paragraph in exported.paragraphs)

            self.assertIn("1 服务范围、服务内容", text)
            self.assertIn("1.1 服务范围", text)
            self.assertIn("本节说明服务范围。", text)
            self.assertNotIn("4 1 服务范围、服务内容", text)
            self.assertNotIn("4.1.1服务范围", text)

    def test_history_case_match_terms_prioritize_specific_project_terms(self):
        from backend.app.services.history_case_service import HistoryCaseService

        terms = HistoryCaseService._extract_match_terms(
            "中国石油辽宁销售公司2025年油库、加油站、新能源工程项目设计服务项目 工程设计 设计方案"
        )

        self.assertIn("中国石油辽宁销售公司2025年油库、加油站、新能源工程项目", terms[0])
        self.assertLess(terms.index("油库 工程设计"), terms.index("油库"))

    def test_history_case_match_deduplicates_same_project_document_hits(self):
        from backend.app.services.history_case_service import HistoryCaseService

        original_tree_root = HistoryCaseService.PAGEINDEX_TREE_ROOT
        with tempfile.TemporaryDirectory() as temp_dir:
            tree_root = Path(temp_dir) / "pageindex_trees"
            tree_payload = {
                "line_count": 2,
                "structure": [
                    {
                        "title": "服务方案",
                        "node_id": "0001",
                        "line_num": 1,
                        "text": "# 服务方案\n油库 加油站 工程设计 服务方案",
                        "nodes": [],
                    }
                ],
            }
            try:
                for relative, doc_name in (
                    ("2025/技术1/p1/技术标-1.json", "重复文档油库设计项目"),
                    ("2025/技术1/p1/技术标-2.json", "重复文档油库设计项目"),
                    ("2025/技术1/p2/技术标.json", "单文档油库设计项目"),
                ):
                    path = tree_root / relative
                    path.parent.mkdir(parents=True, exist_ok=True)
                    payload = {**tree_payload, "doc_name": doc_name}
                    path.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")
                HistoryCaseService.PAGEINDEX_TREE_ROOT = tree_root
                candidates = HistoryCaseService.match_candidates("油库 工程设计", {}, limit=2)
                scores = {item["project_id"]: item["score"] for item in candidates}
                self.assertEqual(scores.get("p1"), scores.get("p2"))
            finally:
                HistoryCaseService.PAGEINDEX_TREE_ROOT = original_tree_root

    def test_history_search_uses_pageindex_nodes_as_primary_source(self):
        from backend.app.services.history_case_service import HistoryCaseService

        original_tree_root = HistoryCaseService.PAGEINDEX_TREE_ROOT
        with tempfile.TemporaryDirectory() as temp_dir:
            tree_root = Path(temp_dir) / "pageindex_trees"
            tree_path = tree_root / "2025/技术1/p1/历史技术标.json"
            tree_path.parent.mkdir(parents=True, exist_ok=True)
            try:
                tree_path.write_text(
                    json.dumps(
                        {
                            "doc_name": "历史油库设计服务项目",
                            "line_count": 10,
                            "structure": [
                                {
                                    "title": "技术标",
                                    "node_id": "0001",
                                    "line_num": 1,
                                    "text": "# 技术标",
                                    "nodes": [
                                        {
                                            "title": "服务范围",
                                            "node_id": "0002",
                                            "line_num": 8,
                                            "text": "## 服务范围\n油库工程设计服务范围包括工艺、建筑、结构、给排水等专业。",
                                            "nodes": [],
                                        }
                                    ],
                                }
                            ],
                        },
                        ensure_ascii=False,
                    ),
                    encoding="utf-8",
                )
                HistoryCaseService.PAGEINDEX_TREE_ROOT = tree_root
                rows = HistoryCaseService.search("油库 工程设计 服务范围", limit=1)
                self.assertTrue(rows[0]["pageindex_node_id"].endswith(":0002"))
                self.assertEqual(rows[0]["node_path"], ["技术标", "服务范围"])
                self.assertIn("油库工程设计", rows[0]["node_text"])

                context = HistoryCaseService.load_pageindex_context_for_candidate({
                    "project_id": "p1",
                    "best_node_text": rows[0]["node_text"],
                    "pageindex_tree_path": rows[0]["pageindex_tree_path"],
                })
                self.assertIn("最相关 PageIndex 节点", context)
                self.assertIn("服务范围", context)
            finally:
                HistoryCaseService.PAGEINDEX_TREE_ROOT = original_tree_root

    def test_history_search_uses_sqlite_recall_with_pageindex_context(self):
        from backend.app.services.history_case_service import HistoryCaseService

        original_tree_root = HistoryCaseService.PAGEINDEX_TREE_ROOT
        original_db_path = HistoryCaseService.HISTORY_DB_PATH
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            tree_root = root / "pageindex_trees"
            tree_path = tree_root / "2025/技术1/p1/历史技术标.json"
            db_path = root / "history_cases.sqlite3"
            tree_path.parent.mkdir(parents=True, exist_ok=True)
            tree_path.write_text(
                json.dumps(
                    {
                        "doc_name": "SQLite召回油库设计服务项目",
                        "structure": [
                            {
                                "title": "技术标",
                                "node_id": "0001",
                                "line_num": 1,
                                "text": "# 技术标",
                                "nodes": [
                                    {
                                        "title": "项目组成人员",
                                        "node_id": "0002",
                                        "line_num": 5,
                                        "text": "## 项目组成人员\n油库工程设计项目组成人员包括项目负责人、工艺、结构专业人员。",
                                        "nodes": [],
                                    }
                                ],
                            }
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            with sqlite3.connect(str(db_path)) as connection:
                connection.executescript(
                    """
                    CREATE TABLE case_projects (
                        id TEXT PRIMARY KEY,
                        year TEXT NOT NULL,
                        batch TEXT NOT NULL,
                        sequence TEXT NOT NULL,
                        result TEXT NOT NULL,
                        subject TEXT NOT NULL,
                        primary_domain TEXT NOT NULL,
                        primary_subdomain TEXT NOT NULL,
                        domain_confidence REAL NOT NULL,
                        domain_keywords TEXT NOT NULL,
                        title TEXT NOT NULL,
                        folder_name TEXT NOT NULL,
                        source_path TEXT NOT NULL,
                        imported_at TEXT NOT NULL
                    );
                    CREATE TABLE case_documents (
                        id TEXT PRIMARY KEY,
                        project_id TEXT NOT NULL,
                        file_name TEXT NOT NULL,
                        extension TEXT NOT NULL,
                        source_path TEXT NOT NULL,
                        markdown_path TEXT NOT NULL,
                        block_json_path TEXT NOT NULL,
                        html_preview_path TEXT NOT NULL,
                        asset_dir TEXT NOT NULL,
                        pageindex_tree_path TEXT NOT NULL,
                        text_chars INTEGER NOT NULL,
                        node_count INTEGER NOT NULL,
                        status TEXT NOT NULL,
                        error TEXT NOT NULL,
                        indexed_at TEXT NOT NULL
                    );
                    CREATE TABLE pageindex_nodes (
                        id TEXT PRIMARY KEY,
                        document_id TEXT NOT NULL,
                        node_id TEXT NOT NULL,
                        parent_node_id TEXT,
                        title TEXT NOT NULL,
                        line_num INTEGER NOT NULL,
                        level INTEGER NOT NULL,
                        text TEXT NOT NULL
                    );
                    """
                )
                connection.execute(
                    "INSERT INTO case_projects VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
                    (
                        "p1", "2025", "技术1", "1", "中标", "华正", "石油", "油库", 0.9,
                        "[]", "SQLite召回油库设计服务项目", "1【中标】SQLite召回油库设计服务项目",
                        str(root / "project"), "2026-01-01T00:00:00Z",
                    ),
                )
                connection.execute(
                    "INSERT INTO case_documents VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
                    (
                        "doc-sqlite", "p1", "SQLite召回油库设计服务项目-2025.docx", ".docx",
                        str(root / "source.docx"), "", "", "", "", str(tree_path), 100, 2,
                        "indexed", "", "2026-01-01T00:00:00Z",
                    ),
                )
                connection.execute(
                    "INSERT INTO pageindex_nodes VALUES (?,?,?,?,?,?,?,?)",
                    (
                        "doc-sqlite:0002", "doc-sqlite", "0002", "0001", "项目组成人员", 5, 2,
                        "SQLite 中只用于召回，完整正文应来自 PageIndex JSON。",
                    ),
                )

            try:
                HistoryCaseService.PAGEINDEX_TREE_ROOT = tree_root
                HistoryCaseService.HISTORY_DB_PATH = db_path
                rows = HistoryCaseService.search("油库 项目组成人员", limit=1)
                self.assertEqual(rows[0]["pageindex_node_id"], "doc-sqlite:0002")
                self.assertEqual(rows[0]["node_path"], ["技术标", "项目组成人员"])
                self.assertIn("油库工程设计项目组成人员", rows[0]["node_text"])
            finally:
                HistoryCaseService.PAGEINDEX_TREE_ROOT = original_tree_root
                HistoryCaseService.HISTORY_DB_PATH = original_db_path

    def test_history_reference_rule_based_profile_is_usable(self):
        from backend.app.routers.history_cases import _build_rule_based_reference_profile

        profile = _build_rule_based_reference_profile(
            "# 技术标\n\n## 服务方案\n\n### 质量保证措施\n正文",
            {"project_title": "历史油库设计案例"},
            "model unavailable",
        )

        self.assertEqual(profile["profile_name"], "历史案例库规则匹配模板")
        self.assertTrue(profile["outline_template"])
        self.assertTrue(profile["chapter_blueprints"])
        self.assertIn("model unavailable", profile["quality_risks"][-1]["fix_rule"])


if __name__ == "__main__":
    unittest.main()
